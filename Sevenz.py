import datetime
import json
import os
import re
import time
from io import BytesIO
from pathlib import Path

import streamlit as st
from fpdf import FPDF
from gtts import gTTS
from openai import OpenAI

try:
    import requests

    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Optional dependency for future voice recording feature
try:
    from audio_recorder_streamlit import audio_recorder

    HAS_AUDIO = True
except ImportError:
    HAS_AUDIO = False


GEMINI_MODEL_CACHE = {}
GEMINI_MODEL_CACHE_TTL_SEC = 900
GEMINI_CONTINUE_ROUNDS = 4


STRUCTURES = {
    "ធម្មតា (Standard)": "[Intro]\n[Verse 1]\n[Pre-Chorus]\n[Chorus]\n[Verse 2]\n[Chorus]\n[Bridge]\n[Chorus]\n[Outro]",
    "រ៉េប (Rap/Hip-Hop)": "[Intro]\n[Verse 1 (Rap)]\n[Chorus (Hook)]\n[Verse 2 (Rap)]\n[Chorus (Hook)]\n[Bridge (Rap)]\n[Chorus (Hook)]\n[Outro]",
    "EDM/Dance (Drop)": "[Intro]\n[Verse 1]\n[Build-up]\n[Drop]\n[Verse 2]\n[Build-up]\n[Drop]\n[Bridge]\n[Outro]",
    "បាឡាត (Ballad)": "[Intro]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Chorus]\n[Bridge]\n[Big Chorus]\n[Outro]",
    "Acoustic Intimate": "[Intro]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Chorus]\n[Outro]",
    "R&B Groove": "[Intro]\n[Verse 1]\n[Pre-Chorus]\n[Chorus]\n[Verse 2]\n[Pre-Chorus]\n[Chorus]\n[Bridge]\n[Final Chorus]\n[Outro]",
    "Cinematic Story": "[Prologue]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Bridge]\n[Climax Chorus]\n[Epilogue]",
    "K-Pop Performance": "[Intro]\n[Verse 1]\n[Pre-Chorus]\n[Chorus]\n[Verse 2]\n[Rap Break]\n[Pre-Chorus]\n[Dance Break]\n[Final Chorus]\n[Outro]",
    "Lo-fi Minimal": "[Intro]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Instrumental]\n[Chorus]\n[Outro]",
    "Rock Anthem": "[Intro Riff]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Chorus]\n[Guitar Solo]\n[Bridge]\n[Final Chorus]\n[Outro]",
    "Afrobeat Pop": "[Intro]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Chorus]\n[Breakdown]\n[Final Chorus]\n[Outro]",
    "Gospel/Uplift": "[Intro]\n[Verse 1]\n[Chorus]\n[Verse 2]\n[Key Change Chorus]\n[Outro]",
}

STYLES = {
    "ជ្រើសរើស...": "",
    "Cinematic (Hans Zimmer)": "Cinematic, Orchestral, Epic, Hans Zimmer style",
    "Pop Modern (The Weeknd)": "Synth-pop, 80s vibes, The Weeknd style",
    "Sad Indie (Lana Del Rey)": "Dream pop, melancholic, vintage, Lana Del Rey style",
    "EDM High Energy (Avicii)": "Progressive House, energetic, uplifting, Avicii style",
    "Khmer Modern": "Modern Khmer pop, emotional, acoustic guitar, soft drums",
    "Chinese Emotional": "Mandopop, emotional, piano, orchestral strings, sad narrative",
    "Neo Soul (H.E.R.)": "Neo soul, warm guitars, intimate vocals, H.E.R.-inspired",
    "Alt Rock (Arctic Monkeys)": "Alternative rock, groovy bass, edgy vocal phrasing",
    "Dream Pop (Cigarettes After Sex)": "Dream pop, airy reverb, soft melancholy",
    "Synthwave (Kavinsky)": "Retro synthwave, neon atmosphere, driving bass",
    "Afrobeats (Wizkid)": "Afrobeats bounce, rhythmic groove, sunny energy",
    "K-Pop Modern": "K-pop polished production, catchy hooks, dynamic sections",
    "Latin Pop (Rosalia)": "Latin pop fusion, bold rhythm, modern vocal textures",
    "Trap Soul": "Trap drums, dark keys, emotional melodic topline",
    "Indie Folk (Bon Iver)": "Indie folk, organic textures, introspective mood",
    "Lo-fi Chill": "Lo-fi beats, vinyl texture, mellow nostalgic vibe",
    "Orchestral Ballad": "Piano and strings, dramatic dynamics, heartfelt storytelling",
    "Khmer Romvong Modern": "Modern Romvong feel, Khmer melodic phrasing, festive rhythm",
    "Khmer Saravan Pop": "Saravan groove with contemporary pop arrangement",
}

GENRES = [
    "Pop",
    "Pop Rock",
    "Indie Pop",
    "R&B",
    "Neo Soul",
    "Hip-Hop",
    "Trap",
    "Drill",
    "Afrobeats",
    "Latin Pop",
    "Reggaeton",
    "Country Pop",
    "Funk",
    "Disco",
    "Synthwave",
    "Deepchill Melodic",
    "Lo-Fi Indie Folk",
    "Indie Folk",
    "Sad Pop",
    "Slowcore",
    "Rock",
    "Alternative Rock",
    "Hard Rock",
    "Punk Rock",
    "Metalcore",
    "EDM / House",
    "Progressive House",
    "Future Bass",
    "Techno",
    "Trance",
    "Drum & Bass",
    "Dubstep",
    "Cinematic / Epic",
    "Soundtrack",
    "Gospel",
    "Jazz",
    "Blues",
    "Khmer Pop",
    "Khmer Ballad",
    "Khmer Hip-Hop",
    "Mandopop",
    "J-Pop",
    "K-Pop",
    "✍️ វាយបញ្ចូលខ្លួនឯង",
]

MOODS = [
    "Relaxing / Chill",
    "Atmospheric",
    "Deep Emotional",
    "Energetic / Hype",
    "Romantic",
    "Epic / Triumphant",
    "Dark / Mysterious",
    "Angry / Aggressive",
    "Hopeful / Uplifting",
    "Melancholic / Nostalgic",
    "Confident / Boss",
    "Dreamy / Floating",
    "Spiritual / Healing",
    "Lonely / Empty",
    "Bittersweet",
    "Festival / Party",
    "Cinematic Tension",
    "Motivational",
    "Heartbreak",
    "Revenge",
    "Prayer / Faith",
    "✍️ វាយបញ្ចូលខ្លួនឯង",
]

INSTS = [
    "Piano",
    "Grand Piano",
    "Upright Piano",
    "Atmospheric Pad",
    "String Ensemble",
    "Solo Violin",
    "Cello",
    "Choir",
    "Flute",
    "Saxophone",
    "Trumpet",
    "Electric Guitar",
    "Acoustic Guitar",
    "Bass Guitar",
    "Ukulele",
    "Synth / Keyboards",
    "Synth Bass",
    "Arpeggiator",
    "Drums / Beats",
    "808 Kit",
    "Boom Bap Drums",
    "Trap Hi-hats",
    "Live Drum Kit",
    "Percussion",
    "Afro Percussion",
    "Lo-Fi Vinyl Crackle",
    "Full Orchestra",
    "Harp",
    "Ambient FX",
    "Gamelan / Ethnic Percussion",
    "✍️ វាយបញ្ចូលខ្លួនឯង",
]

LANG_TTS = {
    "ខ្មែរ": "km",
    "English": "en",
    "ចិន (Chinese)": "zh-CN",
}

LANG_LABEL = {
    "ខ្មែរ": "Khmer",
    "English": "English",
    "ចិន (Chinese)": "Chinese",
}

REWRITE_ACTIONS = {
    "Shorten (ខ្លីជាងមុន)": "Rewrite the lyrics to be about 30% shorter while keeping core meaning and hook.",
    "More Emotional (បន្ថែមអារម្មណ៍)": "Rewrite with deeper emotional impact and vivid imagery.",
    "Cleaner Rhyme (ពាក្យជួនល្អជាងមុន)": "Improve rhyme consistency and line endings while keeping meaning.",
    "Simpler Words (ពាក្យសាមញ្ញ)": "Rewrite using simpler, clearer language for broader audience.",
    "Radio-friendly": "Rewrite to be radio-friendly, catchy, and easy to sing along.",
}

JUDGE_PROFILES = {
    "Strict Commercial": (
        "You are a strict commercial songwriting judge focused on hit potential. "
        "Prioritize catchy hooks, memorability, clear structure, concise wording, and mass singability."
    ),
    "Artistic/Poetic": (
        "You are an artistic songwriting judge focused on poetic depth. "
        "Prioritize imagery, emotional nuance, originality, literary quality, and expressive language."
    ),
    "Balanced": (
        "You are a balanced songwriting judge combining commercial viability and artistic depth. "
        "Weigh hook strength, memorability, singability, emotional nuance, and originality equally."
    ),
}

PROVIDERS = ["OpenAI", "Gemini"]
MODEL_OPTIONS = {
    "OpenAI": [
        "gpt-4.1-mini",
        "gpt-4.1",
        "gpt-4o-mini",
        "gpt-4o",
        "gpt-4.5-preview",
        "o4-mini",
        "o3",
        "o1",
    ],
    "Gemini": ["gemini-2.5-flash", "gemini-2.5-pro-latest"],
}

GEMINI_FALLBACK_MODELS = ["gemini-2.5-flash", "gemini-2.5-pro-latest"]
SETTINGS_PATH = Path(__file__).with_name(".sevenz_settings.json")
ENV_PATH = Path(__file__).with_name(".env")
ENV_KEY_CANDIDATES = {
    "OpenAI": ["OPENAI_API_KEY"],
    "Gemini": ["GEMINI_API_KEY", "GOOGLE_API_KEY"],
    "Suno": ["SUNO_API_KEY"],
}


def load_local_settings():
    try:
        if SETTINGS_PATH.exists():
            return json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return {}


def coerce_gemini_settings(settings):
    coerced = dict(settings or {})
    if coerced.get("api_provider") == "Gemini":
        if coerced.get("api_model_choice") not in ["gemini-2.5-flash", "gemini-2.5-pro-latest"]:
            coerced["api_model_choice"] = "gemini-2.5-flash"
    return coerced


def save_local_settings(data):
    try:
        SETTINGS_PATH.write_text(json.dumps(coerce_gemini_settings(data), ensure_ascii=False, indent=2), encoding="utf-8")
        return True
    except Exception:
        return False


def parse_dotenv_file(dotenv_path):
    values = {}
    try:
        if not dotenv_path.exists():
            return values

        for raw in dotenv_path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            if line.startswith("export "):
                line = line[7:].strip()
            if "=" not in line:
                continue

            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")

            if key:
                values[key] = value
    except Exception:
        return {}

    return values


def resolve_named_api_key(manual_key, names):
    direct_key = (manual_key or "").strip()
    if direct_key:
        return direct_key, "manual"

    try:
        for name in names:
            secret_value = str(st.secrets.get(name, "")).strip()
            if secret_value:
                return secret_value, "secrets"
    except Exception:
        pass

    for name in names:
        env_value = os.getenv(name, "").strip()
        if env_value:
            return env_value, "env"

    if "dotenv_cache" not in st.session_state:
        st.session_state.dotenv_cache = parse_dotenv_file(ENV_PATH)

    dotenv_values = st.session_state.dotenv_cache
    for name in names:
        dot_env_value = str(dotenv_values.get(name, "")).strip()
        if dot_env_value:
            return dot_env_value, "dotenv"

    return "", ""


def resolve_api_key(provider, manual_key):
    return resolve_named_api_key(manual_key, ENV_KEY_CANDIDATES.get(provider, []))


def normalize_url(base_url, path):
    base = (base_url or "").strip().rstrip("/")
    route = (path or "").strip()
    if not route.startswith("/"):
        route = "/" + route
    return base + route


def normalize_bearer_token(token):
    value = (token or "").strip().strip('"').strip("'")
    if value.lower().startswith("bearer "):
        value = value[7:].strip()
    return value


def extract_first_url_from_obj(obj):
    if isinstance(obj, str) and obj.startswith("http"):
        return obj
    if isinstance(obj, list):
        for item in obj:
            found = extract_first_url_from_obj(item)
            if found:
                return found
    if isinstance(obj, dict):
        preferred_keys = ["audio_url", "audioUrl", "url", "stream_url", "download_url", "media_url"]
        for key in preferred_keys:
            value = obj.get(key)
            if isinstance(value, str) and value.startswith("http"):
                return value
        for value in obj.values():
            found = extract_first_url_from_obj(value)
            if found:
                return found
    return ""


def extract_generation_id(obj):
    if isinstance(obj, str):
        return ""
    if isinstance(obj, list):
        for item in obj:
            found = extract_generation_id(item)
            if found:
                return found
    if isinstance(obj, dict):
        for key in ["taskId", "task_id", "id", "generation_id", "job_id", "uuid"]:
            value = obj.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
        for value in obj.values():
            found = extract_generation_id(value)
            if found:
                return found
    return ""


def extract_generation_status(obj):
    if isinstance(obj, str):
        low = obj.lower().strip()
        if low in [
            "pending",
            "queued",
            "processing",
            "running",
            "completed",
            "done",
            "failed",
            "error",
            "success",
            "text_success",
            "first_success",
            "create_task_failed",
            "generate_audio_failed",
            "callback_exception",
            "sensitive_word_error",
        ]:
            return low
        return ""
    if isinstance(obj, list):
        for item in obj:
            found = extract_generation_status(item)
            if found:
                return found
    if isinstance(obj, dict):
        for key in ["status", "state", "task_status", "generation_status"]:
            value = obj.get(key)
            found = extract_generation_status(value)
            if found:
                return found
        for value in obj.values():
            found = extract_generation_status(value)
            if found:
                return found
    return ""


def suno_api_request(api_key, method, url, payload=None, timeout_sec=90):
    if not HAS_REQUESTS:
        return False, "requests library not installed", {}

    clean_key = normalize_bearer_token(api_key)
    if not clean_key:
        return False, "Missing API key. Please provide a valid Suno API key.", {}

    headers = {
        "Authorization": f"Bearer {clean_key}",
        "Accept": "application/json",
        "Content-Type": "application/json",
        "X-API-Key": clean_key,
    }
    try:
        if method == "POST":
            response = requests.post(url, json=payload or {}, headers=headers, timeout=timeout_sec)
        else:
            response = requests.get(url, headers=headers, timeout=timeout_sec)

        if response.status_code == 401:
            return (
                False,
                "HTTP 401 Unauthorized. Check API key validity, account credits/permissions, and ensure Base URL is https://api.sunoapi.org.",
                {"raw": response.text[:1200]},
            )

        if response.status_code >= 400:
            return False, f"HTTP {response.status_code}: {response.text[:500]}", {}

        content_type = (response.headers.get("Content-Type", "") or "").lower()
        body_text = response.text or ""

        # Common misconfiguration: calling suno.com page instead of API endpoint.
        if "text/html" in content_type or body_text.lstrip().startswith("<!doctype html") or body_text.lstrip().startswith("<html"):
            snippet = body_text[:220].replace("\n", " ")
            return (
                False,
                "Endpoint returned HTML (not API JSON). "
                "Please check Base URL/Create Path. Example: use your API host, not https://suno.com homepage. "
                f"Preview: {snippet}",
                {"raw": body_text},
            )

        try:
            data = response.json()
        except Exception:
            return (
                False,
                "Endpoint did not return valid JSON. Please verify API path and credentials.",
                {"raw": body_text},
            )

        return True, "ok", data
    except Exception as exc:
        return False, str(exc), {}


def suno_create_generation(api_key, base_url, create_path, payload):
    url = normalize_url(base_url, create_path)
    ok, msg, data = suno_api_request(api_key, "POST", url, payload=payload)
    if not ok:
        return False, msg, "", data

    generation_id = extract_generation_id(data)
    return True, "created", generation_id, data


def suno_check_generation(api_key, base_url, status_path_template, generation_id):
    path = (status_path_template or "").replace("{id}", generation_id)
    url = normalize_url(base_url, path)
    ok, msg, data = suno_api_request(api_key, "GET", url)
    if not ok:
        return False, msg, "", "", data

    status = extract_generation_status(data)
    audio_url = extract_first_url_from_obj(data)
    return True, "ok", status, audio_url, data


def ensure_daily_usage_state():
    today = datetime.date.today().isoformat()
    if st.session_state.get("usage_day") != today:
        st.session_state.usage_day = today
        st.session_state.usage_requests = 0
        st.session_state.usage_prompt_tokens = 0
        st.session_state.usage_completion_tokens = 0
        st.session_state.usage_total_tokens = 0


def estimate_tokens(text):
    if not text:
        return 0
    return max(1, int(len(text) / 4))


def update_daily_usage(prompt_text, output_text, response):
    ensure_daily_usage_state()
    usage = getattr(response, "usage", None) or getattr(response, "usage_metadata", None)
    if usage is None and isinstance(response, dict):
        usage = response.get("usageMetadata") or response.get("usage_metadata")

    def first_attr(obj, names):
        if not obj:
            return None
        for name in names:
            value = getattr(obj, name, None)
            if value is not None:
                return value
            if isinstance(obj, dict) and name in obj and obj.get(name) is not None:
                return obj.get(name)
        return None

    prompt_tokens = first_attr(usage, ["prompt_tokens", "prompt_token_count", "input_tokens", "input_token_count"])
    completion_tokens = first_attr(usage, ["completion_tokens", "candidates_token_count", "output_tokens", "output_token_count"])
    total_tokens = first_attr(usage, ["total_tokens", "total_token_count"])

    if prompt_tokens is None:
        prompt_tokens = estimate_tokens(prompt_text)
    if completion_tokens is None:
        completion_tokens = estimate_tokens(output_text)
    if total_tokens is None:
        total_tokens = prompt_tokens + completion_tokens

    st.session_state.usage_requests += 1
    st.session_state.usage_prompt_tokens += int(prompt_tokens)
    st.session_state.usage_completion_tokens += int(completion_tokens)
    st.session_state.usage_total_tokens += int(total_tokens)


def update_health_from_result(ok, message, provider, model):
    if ok:
        st.session_state.health_key_status = "Valid"
        st.session_state.health_model_status = f"Available ({model})"
        st.session_state.health_quota_status = "OK"
        st.session_state.health_last_message = "Connection OK"
        return

    msg = (message or "").lower()
    st.session_state.health_key_status = "Unknown"
    st.session_state.health_model_status = f"Check model ({model})"
    st.session_state.health_quota_status = "Unknown"
    st.session_state.health_last_message = message

    if "invalid" in msg or "401" in msg:
        st.session_state.health_key_status = "Invalid"
    if is_quota_error(message):
        st.session_state.health_quota_status = "Exceeded"
    if "no longer available" in msg or "404" in msg:
        st.session_state.health_model_status = "Unavailable"


def get_effective_max_tokens(default_max_tokens):
    if st.session_state.get("cost_safe_mode", False):
        cap = int(st.session_state.get("cost_safe_cap", 500))
        return min(default_max_tokens, cap)
    return default_max_tokens


def is_quota_error(err_text):
    lower = err_text.lower()
    return (
        "error code: 429" in lower
        or "resource_exhausted" in lower
        or "quota exceeded" in lower
        or "rate limit" in lower
    )


def format_provider_error(err_text, provider, model):
    lower = err_text.lower()

    if is_quota_error(err_text):
        retry_match = re.search(r"retry in ([0-9.]+)s", lower)
        retry_note = f" សូមរង់ចាំប្រហែល {retry_match.group(1)}s ហើយសាកម្តងទៀត។" if retry_match else ""
        return (
            f"{provider} quota បានអស់ សម្រាប់ model '{model}'. "
            "API Key របស់អ្នកត្រឹមត្រូវ ប៉ុន្តែគម្រោងបច្ចុប្បន្នមិនមាន quota សល់។"
            f"{retry_note}"
        )

    if "error code: 404" in lower and (
        "no longer available" in lower
        or "not found for api version" in lower
        or "not found" in lower
    ):
        return f"Model '{model}' មិនមាន ឬមិនគាំទ្រសម្រាប់ API version នេះទេ។ សូមប្តូរទៅ model ផ្សេង។"

    if "error code: 401" in lower or "invalid api key" in lower:
        return f"API Key មិនត្រឹមត្រូវសម្រាប់ {provider}។ សូមពិនិត្យ key ម្តងទៀត។"

    return f"{provider} error: {err_text}"


def gemini_generate_content_rest(api_key, model, system_prompt, user_prompt, temperature, max_tokens):
    api_versions = ["v1beta", "v1"]
    last_error = None

    for api_version in api_versions:
        url = f"https://generativelanguage.googleapis.com/{api_version}/models/{model}:generateContent"
        headers = {
            "Content-Type": "application/json; charset=utf-8",
        }
        params = {"key": api_key}
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": user_prompt}],
                }
            ],
            "systemInstruction": {
                "parts": [{"text": system_prompt}],
            },
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
            },
        }
        response = requests.post(url, headers=headers, params=params, json=payload, timeout=90)
        if response.status_code >= 400:
            try:
                error_json = response.json()
                error_message = error_json.get("error", {}).get("message", response.text)
            except Exception:
                error_message = response.text
            last_error = RuntimeError(f"HTTP {response.status_code}: {error_message}")
            if response.status_code == 404:
                continue
            raise last_error
        return response.json()

    if last_error:
        raise last_error
    raise RuntimeError("Gemini request failed without a response.")


def gemini_list_models_rest(api_key):
    now_ts = time.time()
    cached = GEMINI_MODEL_CACHE.get(api_key)
    if cached and (now_ts - cached.get("ts", 0) <= GEMINI_MODEL_CACHE_TTL_SEC):
        return list(cached.get("models", []))

    for api_version in ["v1beta", "v1"]:
        url = f"https://generativelanguage.googleapis.com/{api_version}/models"
        response = requests.get(url, params={"key": api_key}, timeout=30)
        if response.status_code >= 400:
            continue
        try:
            data = response.json()
            models = data.get("models", []) if isinstance(data, dict) else []
            model_names = []
            for item in models:
                if isinstance(item, dict):
                    name = item.get("name", "")
                    if name.startswith("models/"):
                        name = name.split("models/", 1)[1]
                    if name:
                        model_names.append(name)
            GEMINI_MODEL_CACHE[api_key] = {"ts": now_ts, "models": model_names}
            return model_names
        except Exception:
            return []

    if cached:
        # Fallback to stale cache when model listing endpoint is temporarily unavailable.
        return list(cached.get("models", []))
    return []


def gemini_candidate_models(api_key, requested_model):
    requested = (requested_model or "").strip()
    discovered = gemini_list_models_rest(api_key)

    preferred = []
    if requested:
        preferred.append(requested)

    flash_models = [m for m in discovered if "flash" in m.lower()]
    pro_models = [m for m in discovered if "pro" in m.lower()]

    for item in flash_models + pro_models + GEMINI_FALLBACK_MODELS:
        if item and item not in preferred:
            preferred.append(item)

    if not preferred:
        preferred = ["gemini-2.5-flash", "gemini-2.5-pro-latest"]

    return preferred


def gemini_extract_text(response_json):
    if not isinstance(response_json, dict):
        return "", "Gemini response was not JSON"

    direct_text = response_json.get("text")
    if isinstance(direct_text, str) and direct_text.strip():
        return direct_text.strip(), ""

    candidates = response_json.get("candidates") or []
    if candidates:
        best_text = ""
        finish_reasons = []
        for candidate in candidates:
            if not isinstance(candidate, dict):
                continue
            content = candidate.get("content") or {}
            parts = content.get("parts") or []
            parts_text = []
            for part in parts:
                if isinstance(part, dict):
                    if isinstance(part.get("text"), str) and part.get("text").strip():
                        parts_text.append(part.get("text").strip())
                    elif isinstance(part.get("inline_data"), dict):
                        parts_text.append("[inline_data]")
                    elif isinstance(part.get("functionCall"), dict):
                        parts_text.append("[functionCall]")
            finish_reason = candidate.get("finishReason") or candidate.get("finish_reason") or ""
            if finish_reason:
                finish_reasons.append(str(finish_reason))

            text = "\n".join(parts_text).strip()
            if text and not best_text:
                best_text = text

        if best_text:
            if finish_reasons:
                unique_reasons = []
                for reason in finish_reasons:
                    if reason not in unique_reasons:
                        unique_reasons.append(reason)
                return best_text, f"finishReason={','.join(unique_reasons)}"
            return best_text, ""

        if finish_reasons:
            unique_reasons = []
            for reason in finish_reasons:
                if reason not in unique_reasons:
                    unique_reasons.append(reason)
            return "", f"finishReason={','.join(unique_reasons)}"

    prompt_feedback = response_json.get("promptFeedback") or response_json.get("prompt_feedback") or {}
    if isinstance(prompt_feedback, dict):
        block_reason = prompt_feedback.get("blockReason") or prompt_feedback.get("block_reason") or ""
        safety_ratings = prompt_feedback.get("safetyRatings") or prompt_feedback.get("safety_ratings") or []
        if block_reason:
            return "", f"prompt blocked: {block_reason}"
        if safety_ratings:
            return "", "prompt blocked by safety settings"

    return "", "no text returned"


# -------------------- AI Helpers --------------------
def call_ai(
    api_key,
    system_prompt,
    user_prompt,
    provider="OpenAI",
    model="gpt-4.1-mini",
    temperature=0.8,
    max_tokens=1200,
    track_usage=True,
):
    try:
        retry_delays = [0, 2, 5, 10]

        def do_request(client, request_model):
            return client.chat.completions.create(
                model=request_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )

        if provider == "Gemini":
            candidate_models = gemini_candidate_models(api_key, model)
            last_error = ""

            def continue_gemini_output(chosen_model, base_user_prompt, existing_text, request_max_tokens):
                continuation_prompt = (
                    "Continue the same answer from where it stopped.\n"
                    "Do not repeat earlier lines.\n"
                    "Return only the continuation text.\n\n"
                    f"Original request:\n{base_user_prompt}\n\n"
                    "Existing output tail:\n"
                    f"{existing_text[-3000:]}\n\n"
                    "Continue now:"
                )
                response_json = gemini_generate_content_rest(
                    api_key=api_key,
                    model=chosen_model,
                    system_prompt=system_prompt,
                    user_prompt=continuation_prompt,
                    temperature=temperature,
                    max_tokens=request_max_tokens,
                )
                return gemini_extract_text(response_json), response_json

            for chosen_model in candidate_models:
                # Progressively increase output budget for Gemini until the API ceiling.
                token_budgets = []
                current_budget = max(64, int(max_tokens))
                while True:
                    if current_budget not in token_budgets:
                        token_budgets.append(current_budget)
                    if current_budget >= 4096:
                        break
                    current_budget = min(4096, current_budget * 2)

                for current_max_tokens in token_budgets:
                    should_try_next_budget = False
                    for delay in retry_delays:
                        if delay > 0:
                            time.sleep(delay)
                        try:
                            response_json = gemini_generate_content_rest(
                                api_key=api_key,
                                model=chosen_model,
                                system_prompt=system_prompt,
                                user_prompt=user_prompt,
                                temperature=temperature,
                                max_tokens=current_max_tokens,
                            )
                            text, parse_note = gemini_extract_text(response_json)
                            response = response_json
                            if not text.strip():
                                if "finishReason=STOP" in (parse_note or ""):
                                    # Some Gemini responses return STOP with empty first candidate.
                                    # Retry delays/model fallbacks before surfacing an error.
                                    if delay != retry_delays[-1]:
                                        continue
                                if "MAX_TOKENS" in (parse_note or "") and current_max_tokens < token_budgets[-1]:
                                    # Move to the next token budget immediately.
                                    should_try_next_budget = True
                                    break
                                details = parse_note or json.dumps(response_json, ensure_ascii=False)[:800]
                                return False, f"Gemini returned no text ({details})."

                            if "MAX_TOKENS" in (parse_note or ""):
                                assembled_parts = [text.strip()]
                                continuation_source = text
                                for _ in range(GEMINI_CONTINUE_ROUNDS):
                                    (cont_text, cont_note), cont_response_json = continue_gemini_output(
                                        chosen_model,
                                        user_prompt,
                                        continuation_source,
                                        current_max_tokens,
                                    )
                                    if cont_text.strip():
                                        assembled_parts.append(cont_text.strip())
                                        continuation_source = "\n\n".join(assembled_parts)
                                        response = cont_response_json
                                    if "MAX_TOKENS" not in (cont_note or ""):
                                        break
                                text = "\n\n".join([part for part in assembled_parts if part]).strip()

                            if track_usage:
                                update_daily_usage(system_prompt + "\n" + user_prompt, text, response)
                            if chosen_model != model:
                                return True, f"⚠️ Model '{model}' unavailable, switched to '{chosen_model}'.\n\n{text}"
                            return True, text
                        except Exception as exc:
                            err_text = str(exc)
                            last_error = err_text
                            if "404" in err_text and "no longer available" in err_text.lower():
                                break
                            if is_quota_error(err_text):
                                if delay != retry_delays[-1]:
                                    continue
                                return False, format_provider_error(err_text, provider, chosen_model)
                            break
                    if should_try_next_budget:
                        continue
            return False, format_provider_error(last_error, provider, model)
        else:
            client = OpenAI(api_key=api_key)
            last_error = ""
            for delay in retry_delays:
                if delay > 0:
                    time.sleep(delay)
                try:
                    response = do_request(client, model)
                    text = response.choices[0].message.content or ""
                    if not text.strip():
                        return False, "Error: Empty response from AI model."
                    if track_usage:
                        update_daily_usage(system_prompt + "\n" + user_prompt, text, response)
                    return True, text
                except Exception as exc:
                    err_text = str(exc)
                    last_error = err_text
                    if is_quota_error(err_text) and delay != retry_delays[-1]:
                        continue
                    break
            return False, format_provider_error(last_error, provider, model)
    except Exception as exc:
        return False, format_provider_error(str(exc), provider, model)


def test_api_key(api_key, provider, model):
    return call_ai(
        api_key=api_key,
        system_prompt="You are a concise assistant.",
        user_prompt="Reply with exactly API_OK and nothing else.",
        provider=provider,
        model=model,
        temperature=0,
        max_tokens=20,
        track_usage=False,
    )


def is_api_test_success_response(text):
    if not isinstance(text, str) or not text.strip():
        return False
    # Accept API_OK even when the model adds markdown/punctuation.
    return re.search(r"\bAPI\s*[_\-]?\s*OK\b", text, flags=re.IGNORECASE) is not None


def build_song_prompt(lang, title, desc, genre, mood, instruments, structure, style_inf, bpm):
    return f"""Write a complete song with strong emotion and cohesive storytelling.

Language: {LANG_LABEL.get(lang, lang)}
Title: {title or 'Untitled'}
Story: {desc or 'No story provided'}
Genre: {genre}
Mood: {mood}
Instruments: {instruments}
BPM: {bpm}
Style Influence: {style_inf}

Structure:
{structure}

Rules:
- Output only lyrics with section headers.
- Keep it singable and natural.
- Preserve emotional arc from start to finish.
"""


def build_style_summary(genre, mood, instruments, style_inf, bpm):
    return (
        f"Genre: {genre}\n"
        f"Mood: {mood}\n"
        f"Instruments: {instruments}\n"
        f"Style Influence: {style_inf}\n"
        f"BPM: {bpm}"
    )


def push_history(content, source, provider_name, model_name):
    if not content or not content.strip():
        return
    item = {
        "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "source": source,
        "provider": provider_name,
        "model": model_name,
        "content": content,
    }
    st.session_state.lyrics_history.insert(0, item)
    st.session_state.lyrics_history = st.session_state.lyrics_history[:30]


def smart_rewrite(api_key, text, rewrite_instruction, provider, model):
    system_prompt = "You are an expert lyric editor. Keep structure and musical flow natural."
    user_prompt = f"""
Task: {rewrite_instruction}

Lyrics:
{text}

Return only rewritten lyrics.
"""
    return call_ai(
        api_key=api_key,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        provider=provider,
        model=model,
        temperature=0.7,
        max_tokens=get_effective_max_tokens(1200),
    )


def analyze_lyrics(api_key, text, provider, model):
    system_prompt = (
        "You are a professional lyric critic. Score and explain briefly. "
        "Return in clean markdown with bullets."
    )
    user_prompt = f"""
Evaluate the following lyrics.

Lyrics:
{text}

Return:
- Emotional Depth (0-10)
- Rhyme Quality (0-10)
- Flow/Singability (0-10)
- Memorability (0-10)
- 3 specific improvement suggestions
"""
    return call_ai(
        api_key,
        system_prompt,
        user_prompt,
        provider=provider,
        model=model,
        temperature=0.4,
        max_tokens=get_effective_max_tokens(700),
    )


def translate_with_rhyme(api_key, text, target_lang, provider, model):
    system_prompt = (
        "You are an expert lyric translator. Keep the meaning and emotional tone, "
        "and preserve rhyme/musicality as much as possible."
    )
    user_prompt = f"""
Translate these lyrics to {LANG_LABEL.get(target_lang, target_lang)}.
Preserve rhyme feeling and rhythm.

Lyrics:
{text}
"""
    return call_ai(
        api_key,
        system_prompt,
        user_prompt,
        provider=provider,
        model=model,
        temperature=0.7,
        max_tokens=get_effective_max_tokens(1200),
    )


def judge_lyrics_versions(
    api_key,
    version_a,
    version_b,
    version_c,
    provider,
    model,
    judge_profile="Strict Commercial",
):
    system_prompt = JUDGE_PROFILES.get(judge_profile, JUDGE_PROFILES["Strict Commercial"])
    user_prompt = f"""
Evaluate these 3 lyric versions and choose exactly one winner.
Judge profile: {judge_profile}

Version A:
{version_a}

Version B:
{version_b}

Version C:
{version_c}

Return in exactly this format:
WINNER: A or B or C
HOOK: A=x, B=x, C=x (0-10)
EMOTION: A=x, B=x, C=x (0-10)
RHYME: A=x, B=x, C=x (0-10)
SINGABILITY: A=x, B=x, C=x (0-10)
WHY: short explanation in 3-5 lines
IMPROVEMENT: one actionable tip for each version
"""
    return call_ai(
        api_key,
        system_prompt,
        user_prompt,
        provider=provider,
        model=model,
        temperature=0.2,
        max_tokens=get_effective_max_tokens(900),
    )


def is_section_header(line):
    return bool(re.match(r"^\s*\[[^\]]+\]\s*$", line or ""))


def parse_lyric_sections(text):
    sections = []
    current_header = "[Body]"
    buffer = []

    for raw in (text or "").splitlines():
        line = raw.rstrip("\n")
        if is_section_header(line):
            if buffer:
                sections.append((current_header, "\n".join(buffer).strip()))
                buffer = []
            current_header = line.strip()
        else:
            buffer.append(line)

    if buffer:
        sections.append((current_header, "\n".join(buffer).strip()))

    if not sections and (text or "").strip():
        sections.append(("[Body]", text.strip()))

    return sections


def strip_to_words(line):
    cleaned = re.sub(r"[^A-Za-z0-9\s'-]", " ", line or "")
    return [w for w in cleaned.split() if w.strip()]


def estimate_syllables_word(word):
    w = re.sub(r"[^a-z]", "", (word or "").lower())
    if not w:
        return 1
    vowels = "aeiouy"
    count = 0
    prev_vowel = False
    for ch in w:
        is_vowel = ch in vowels
        if is_vowel and not prev_vowel:
            count += 1
        prev_vowel = is_vowel
    if w.endswith("e") and count > 1:
        count -= 1
    return max(1, count)


def estimate_line_syllables(line):
    words = strip_to_words(line)
    if not words:
        return 0
    return sum(estimate_syllables_word(w) for w in words)


def ending_sound_key(line):
    words = strip_to_words(line)
    if not words:
        return ""
    last = re.sub(r"[^a-z0-9]", "", words[-1].lower())
    if not last:
        return ""
    return last[-3:] if len(last) >= 3 else last


def hook_and_singability_report(text):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip() and not is_section_header(ln)]
    if not lines:
        return {
            "hook_score": 0,
            "singability_score": 0,
            "avg_words": 0,
            "avg_syllables": 0,
            "suggestions": ["សូមបញ្ចូលទំនុកច្រៀងជាមុនសិន"],
        }

    normalized_lines = [" ".join(strip_to_words(ln)).lower() for ln in lines]
    words_per_line = [len(strip_to_words(ln)) for ln in lines]
    syllables_per_line = [estimate_line_syllables(ln) for ln in lines]

    freq = {}
    for nl in normalized_lines:
        if len(nl) >= 6:
            freq[nl] = freq.get(nl, 0) + 1
    max_repeat = max(freq.values()) if freq else 1

    avg_words = sum(words_per_line) / len(words_per_line)
    avg_syllables = sum(syllables_per_line) / len(syllables_per_line)
    variance_words = sum((x - avg_words) ** 2 for x in words_per_line) / len(words_per_line)
    rhythm_stability = max(0, 25 - int(variance_words * 1.8))
    brevity = max(0, 25 - int(abs(avg_words - 7) * 4))
    repeat_power = min(30, (max_repeat - 1) * 10)
    vocab = set(" ".join(normalized_lines).split())
    unique_ratio = (len(vocab) / max(1, sum(words_per_line)))
    clarity = min(20, int(unique_ratio * 30))
    hook_score = max(0, min(100, brevity + rhythm_stability + repeat_power + clarity))

    long_lines = sum(1 for n in words_per_line if n > 12)
    short_lines = sum(1 for n in words_per_line if n < 3)
    syllable_range = (max(syllables_per_line) - min(syllables_per_line)) if syllables_per_line else 0
    singability_score = max(
        0,
        min(
            100,
            100 - long_lines * 6 - short_lines * 3 - int(syllable_range * 1.5) - int(variance_words * 2),
        ),
    )

    suggestions = []
    if long_lines > 0:
        suggestions.append("មានបន្ទាត់វែងពេក, ព្យាយាមបំបែក line ឲ្យខ្លីជាងមុន")
    if max_repeat < 2:
        suggestions.append("បន្ថែម Hook line ដែលធ្វើឡើងវិញ 2-3 ដង ដើម្បីចងចាំងាយ")
    if syllable_range > 8:
        suggestions.append("សម្រួលព្យាង្គក្នុងបន្ទាត់ឲ្យជិតគ្នា ដើម្បីងាយច្រៀងលើ beat")
    if hook_score >= 75 and singability_score >= 75:
        suggestions.append("ល្អហើយ: hook និង flow មានតុល្យភាពល្អ")

    return {
        "hook_score": hook_score,
        "singability_score": singability_score,
        "avg_words": round(avg_words, 2),
        "avg_syllables": round(avg_syllables, 2),
        "suggestions": suggestions[:4],
    }


def rhyme_engine_report(text):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip() and not is_section_header(ln)]
    if len(lines) < 2:
        return {
            "rhyme_score": 0,
            "scheme": "N/A",
            "pairs": [],
            "suggestions": ["ត្រូវការបន្ទាត់យ៉ាងហោចណាស់ 2 ដើម្បីវិភាគ rhyme"],
        }

    keys = [ending_sound_key(ln) for ln in lines]
    buckets = {}
    for idx, key in enumerate(keys):
        if key:
            buckets.setdefault(key, []).append(idx)

    paired = sum(len(v) for v in buckets.values() if len(v) >= 2)
    rhyme_ratio = paired / max(1, len(lines))
    rhyme_score = max(0, min(100, int(rhyme_ratio * 100)))

    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    map_key_to_letter = {}
    next_letter = 0
    scheme = []
    for key in keys:
        if not key:
            scheme.append("-")
            continue
        if key not in map_key_to_letter:
            map_key_to_letter[key] = alphabet[next_letter % len(alphabet)]
            next_letter += 1
        scheme.append(map_key_to_letter[key])

    pairs = []
    for key, indexes in buckets.items():
        if len(indexes) >= 2:
            pairs.append(f"-'{key}': lines {', '.join(str(i + 1) for i in indexes[:5])}")

    suggestions = []
    if rhyme_score < 40:
        suggestions.append("ព្យាយាមឲ្យចុងបន្ទាត់ជាប់ៗគ្នា មានសំឡេងចុងដូចគ្នាច្រើនជាងមុន")
    if rhyme_score >= 40 and rhyme_score < 70:
        suggestions.append("អាចបន្ថែម internal rhyme ក្នុង line ដើម្បីឲ្យទំនុកលាន់ស្រួល")
    if rhyme_score >= 70:
        suggestions.append("Rhyme flow ល្អហើយ, អាចបន្ថែមពាក្យរូបភាពដើម្បីឲ្យមានអារម្មណ៍ជ្រាលជ្រៅ")

    return {
        "rhyme_score": rhyme_score,
        "scheme": " ".join(scheme),
        "pairs": pairs[:6],
        "suggestions": suggestions,
    }


def render_score_bar(label, score):
    normalized = max(0, min(100, int(score)))
    left, right = st.columns([1.6, 6])
    with left:
        st.caption(f"{label}: {normalized}")
    with right:
        st.progress(normalized)


def rewrite_section_with_ai(api_key, full_text, section_header, rewrite_instruction, provider, model):
    sections = parse_lyric_sections(full_text)
    if not sections:
        return False, "No lyrics found."

    section_header = (section_header or "").strip()
    if section_header == "Full Lyrics":
        return smart_rewrite(api_key, full_text, rewrite_instruction, provider, model)

    target_index = None
    if section_header == "Auto (First Chorus/Hook)":
        for idx, (hdr, _) in enumerate(sections):
            if "chorus" in hdr.lower() or "hook" in hdr.lower():
                target_index = idx
                break
    else:
        for idx, (hdr, _) in enumerate(sections):
            if hdr == section_header:
                target_index = idx
                break

    if target_index is None:
        return False, "រកមិនឃើញ section ដែលបានជ្រើស។"

    target_header, target_text = sections[target_index]
    system_prompt = "You are an expert lyric editor. Rewrite only the requested section while preserving song identity."
    user_prompt = f"""
Rewrite instruction: {rewrite_instruction}

Target section header: {target_header}

Target section lyrics:
{target_text}

Rules:
- Keep this section singable and emotionally strong.
- Keep section purpose intact (e.g., chorus remains catchy).
- Return only rewritten section body (no extra notes, no markdown).
"""
    ok, rewritten_section = call_ai(
        api_key,
        system_prompt,
        user_prompt,
        provider=provider,
        model=model,
        temperature=0.72,
        max_tokens=get_effective_max_tokens(700),
    )
    if not ok:
        return False, rewritten_section

    new_sections = sections[:]
    new_sections[target_index] = (target_header, rewritten_section.strip())

    rebuilt = []
    for header, body in new_sections:
        rebuilt.append(header)
        rebuilt.append(body.strip())
    final_text = "\n\n".join(chunk for chunk in rebuilt if chunk.strip())
    return True, final_text


# -------------------- Export Helpers --------------------
def make_pdf_bytes(title, body_text, provider_name, model_name):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    def safe_multi_line(text, line_h=8):
        # Normalize to core-font-safe text and avoid tabs causing odd width issues.
        safe = (text or "").replace("\t", "    ").encode("latin-1", errors="replace").decode("latin-1")
        for raw_line in safe.splitlines() or [""]:
            line = raw_line
            # Split very long unbroken tokens so fpdf always has wrap opportunities.
            if len(line) > 140 and " " not in line:
                chunks = [line[i : i + 120] for i in range(0, len(line), 120)]
            else:
                chunks = [line]

            for chunk in chunks:
                try:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(usable_w, line_h, chunk)
                except Exception:
                    # Last-resort fallback: truncate to avoid crashing app export flow.
                    fallback = (chunk[:150] + "...") if len(chunk) > 150 else chunk
                    pdf.set_x(pdf.l_margin)
                    pdf.cell(usable_w, line_h, fallback)
                    pdf.ln(line_h)

    pdf.set_font("Arial", "B", 14)
    safe_multi_line(title, line_h=10)
    pdf.set_font("Arial", size=10)
    safe_multi_line(f"Provider: {provider_name} | Model: {model_name}", line_h=8)
    safe_multi_line(f"Exported at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", line_h=8)
    pdf.ln(2)
    pdf.set_font("Arial", size=11)

    safe_multi_line(body_text, line_h=8)

    pdf_bytes = pdf.output(dest="S").encode("latin-1")
    return BytesIO(pdf_bytes)


def make_docx_bytes(title, body_text, provider_name, model_name):
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Provider: {provider_name}")
    doc.add_paragraph(f"Model: {model_name}")
    doc.add_paragraph(f"Exported at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    doc.add_paragraph(body_text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# -------------------- UI --------------------
st.set_page_config(page_title="SEVENZ Studio Premium", page_icon="🎵", layout="wide")
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;700&family=Noto+Sans+Khmer:wght@400;500;700&display=swap');
:root {
    --bg-0: #0b1119;
    --bg-1: #0f1823;
    --bg-card: rgba(19, 34, 51, 0.44);
    --stroke: rgba(98, 165, 255, 0.24);
    --text-main: #e7f0ff;
    --text-soft: #b8c6da;
    --accent: #37d29e;
    --accent-2: #56a7ff;
    --warn: #ffcc66;
}
html, body, .stApp {
    font-family: 'Noto Sans Khmer', 'Space Grotesk', sans-serif !important;
    font-size: 17px !important;
    color: var(--text-main);
}

/* Keep Streamlit icon fonts intact (prevents labels like "arrow_down" from showing as text). */
span.material-symbols-rounded,
span.material-icons,
[data-testid="stExpanderToggleIcon"] span {
    font-family: "Material Symbols Rounded", "Material Icons" !important;
}

[data-testid="stExpander"] details summary {
    display: flex;
    align-items: center;
    gap: 0.45rem;
    line-height: 1.35;
}

header[data-testid="stHeader"] {
    display: none;
}

.stApp {
    margin-top: 0;
    background:
        radial-gradient(circle at 12% 8%, rgba(86, 167, 255, 0.18), transparent 28%),
        radial-gradient(circle at 88% 18%, rgba(55, 210, 158, 0.17), transparent 28%),
        linear-gradient(135deg, var(--bg-0), #081018 48%, #0d1620);
}

.block-container {
    max-width: 100% !important;
    padding-top: 0.35rem;
    padding-left: 1rem;
    padding-right: 1rem;
    background: rgba(8, 17, 27, 0.19);
    border: 1px solid rgba(128, 178, 238, 0.14);
    border-radius: 18px;
    backdrop-filter: blur(8px);
    -webkit-backdrop-filter: blur(8px);
}

.hero {
    background:
        radial-gradient(circle at 82% 18%, rgba(86, 167, 255, 0.32) 0%, rgba(86, 167, 255, 0.02) 42%),
        radial-gradient(circle at 20% 30%, rgba(55, 210, 158, 0.22) 0%, rgba(55, 210, 158, 0.02) 40%),
        linear-gradient(120deg, rgba(16, 34, 53, 0.61) 0%, rgba(19, 31, 47, 0.58) 55%, rgba(15, 25, 40, 0.61) 100%);
    border: 1px solid var(--stroke);
    border-radius: 18px;
    padding: 16px 18px;
    margin-bottom: 14px;
    width: 100%;
    box-sizing: border-box;
    overflow: hidden;
    backdrop-filter: blur(12px);
    -webkit-backdrop-filter: blur(12px);
}

.hero-title {
    font-size: 25px;
    font-weight: 700;
    color: #f2f7ff;
    font-family: 'Space Grotesk', 'Noto Sans Khmer', sans-serif;
    letter-spacing: 0.2px;
}

.hero-brand {
    display: flex;
    align-items: center;
    gap: 10px;
}

.logo-link {
    text-decoration: none;
    display: inline-block;
    cursor: pointer;
}

.logo-link .logo-7code {
    transition: transform 0.16s ease, box-shadow 0.2s ease, filter 0.2s ease;
}

.logo-link:hover .logo-7code,
.logo-link:focus-visible .logo-7code {
    transform: translateY(-1px) scale(1.03);
    box-shadow: 0 0 0 1px rgba(177, 220, 255, 0.35), 0 8px 18px rgba(41, 120, 204, 0.26);
    filter: saturate(1.12);
}

.logo-link:focus-visible {
    outline: 2px solid rgba(137, 200, 255, 0.75);
    outline-offset: 2px;
    border-radius: 999px;
}

.logo-7code {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    min-width: 82px;
    height: 30px;
    padding: 0 10px;
    border-radius: 999px;
    background: linear-gradient(135deg, rgba(55, 210, 158, 0.32), rgba(86, 167, 255, 0.35));
    border: 1px solid rgba(184, 214, 255, 0.35);
    color: #eaf4ff;
    font-size: 12px;
    font-weight: 700;
    letter-spacing: 0.8px;
    text-transform: uppercase;
    position: relative;
    backdrop-filter: blur(5px);
    -webkit-backdrop-filter: blur(5px);
    animation: logoBounce 2.1s ease-in-out infinite, codeBlurGlow 2.1s ease-in-out infinite;
}

.logo-7code::before {
    content: "";
    position: absolute;
    inset: -6px;
    border-radius: 999px;
    background: radial-gradient(circle, rgba(120, 200, 255, 0.26) 0%, rgba(120, 200, 255, 0) 72%);
    filter: blur(7px);
    z-index: -1;
    animation: logoAura 2.1s ease-in-out infinite;
}

.hero-sub {
    color: #c7d7ed;
    font-size: 14px;
}

.stTextArea textarea, .stTextInput input, .stSelectbox select {
    background: rgba(13, 24, 38, 0.38) !important;
    border-radius: 12px !important;
    border: 1px solid rgba(119, 160, 214, 0.35) !important;
    backdrop-filter: blur(6px);
    -webkit-backdrop-filter: blur(6px);
}

.stButton > button {
    background: linear-gradient(135deg, #2fcb97, #3a8fff) !important;
    color: white !important;
    border-radius: 12px !important;
    border: none !important;
    padding: 0.55rem 1.1rem !important;
    transition: transform 0.14s ease, box-shadow 0.2s ease, filter 0.2s ease !important;
    box-shadow: 0 6px 18px rgba(13, 24, 38, 0.35);
}

.stButton > button:hover {
    transform: translateY(-2px) scale(1.01);
    filter: saturate(1.08);
    box-shadow: 0 8px 20px rgba(10, 20, 34, 0.38);
}

.stButton > button:active {
    transform: translateY(0) scale(0.985);
    box-shadow: 0 4px 12px rgba(10, 20, 34, 0.34);
}

.stButton > button:focus-visible {
    outline: 2px solid rgba(125, 196, 255, 0.75);
    outline-offset: 2px;
}

@keyframes buttonPulse {
    0% {
        box-shadow: 0 6px 16px rgba(13, 24, 38, 0.34);
    }
    50% {
        box-shadow: 0 10px 22px rgba(40, 140, 230, 0.26);
    }
    100% {
        box-shadow: 0 6px 16px rgba(13, 24, 38, 0.34);
    }
}

@media (prefers-reduced-motion: reduce) {
    .stButton > button {
        transition: none !important;
    }
}

@media (max-width: 768px) {
    .block-container {
        padding-left: 0.5rem;
        padding-right: 0.5rem;
        border-radius: 12px;
    }
    .hero {
        padding: 12px 12px;
    }
    .hero-title {
        font-size: 21px;
    }
    .hero-sub {
        font-size: 13px;
    }
    .logo-7code {
        min-width: 72px;
        height: 27px;
        font-size: 11px;
        animation: none;
    }
    .quickbar {
        top: 0.2rem;
        padding: 6px 8px;
    }
}

[data-testid="stTabs"] [role="tablist"] {
    gap: 8px;
    background: rgba(8, 18, 30, 0.32);
    border: 1px solid var(--stroke);
    border-radius: 14px;
    padding: 6px;
    backdrop-filter: blur(10px);
    -webkit-backdrop-filter: blur(10px);
}

[data-testid="stTabs"] [role="tab"] {
    border-radius: 10px;
    padding: 8px 14px;
    color: var(--text-soft);
}

[data-testid="stTabs"] [aria-selected="true"] {
    background: linear-gradient(120deg, rgba(55, 210, 158, 0.22), rgba(86, 167, 255, 0.24));
    color: var(--text-main);
}

.pro-card {
    background: linear-gradient(165deg, rgba(19, 34, 51, 0.46), rgba(14, 27, 42, 0.42));
    border: 1px solid var(--stroke);
    border-radius: 14px;
    padding: 10px 12px;
    margin-bottom: 10px;
    box-shadow: 0 8px 24px rgba(5, 10, 20, 0.22);
    animation: fadeInUp 0.24s ease;
    backdrop-filter: blur(14px);
    -webkit-backdrop-filter: blur(14px);
}

.card-kicker {
    font-size: 12px;
    color: #8db4e6;
    text-transform: uppercase;
    letter-spacing: 0.65px;
    margin-bottom: 4px;
}

.pill {
    display: inline-block;
    font-size: 12px;
    border-radius: 999px;
    padding: 2px 9px;
    border: 1px solid rgba(141, 180, 230, 0.35);
    color: #cfe2fb;
    background: rgba(86, 167, 255, 0.15);
}

.quickbar {
    position: sticky;
    top: 0.5rem;
    z-index: 99;
    background: rgba(7, 13, 22, 0.84);
    border: 1px solid var(--stroke);
    border-radius: 12px;
    padding: 8px 10px;
    margin-bottom: 12px;
    backdrop-filter: blur(8px);
}

@keyframes fadeInUp {
    from {
        opacity: 0;
        transform: translateY(6px);
    }
    to {
        opacity: 1;
        transform: translateY(0px);
    }
}

@keyframes logoBounce {
    0%, 100% {
        transform: translateY(0);
    }
    50% {
        transform: translateY(-5px);
    }
}

@keyframes codeBlurGlow {
    0%, 100% {
        text-shadow: 0 0 0 rgba(173, 230, 255, 0.0);
    }
    50% {
        text-shadow: 0 0 8px rgba(173, 230, 255, 0.6), 0 0 16px rgba(86, 167, 255, 0.35);
    }
}

@keyframes logoAura {
    0%, 100% {
        opacity: 0.28;
        transform: scale(0.98);
    }
    50% {
        opacity: 0.58;
        transform: scale(1.04);
    }
}

.status-good {
    color: #6bf1a4;
    font-weight: 600;
}

.status-warn {
    color: #ffd786;
    font-weight: 600;
}
</style>
""",
    unsafe_allow_html=True,
)

if "notepad_content" not in st.session_state:
    st.session_state.notepad_content = ""
if "last_lyrics" not in st.session_state:
    st.session_state.last_lyrics = ""
if "selected_lang" not in st.session_state:
    st.session_state.selected_lang = "ខ្មែរ"
if "api_provider" not in st.session_state:
    st.session_state.api_provider = "Gemini"
if "api_model_choice" not in st.session_state:
    st.session_state.api_model_choice = "gemini-2.5-flash"
if "api_model_custom" not in st.session_state:
    st.session_state.api_model_custom = ""
if "api_test_message" not in st.session_state:
    st.session_state.api_test_message = ""
if "cost_safe_mode" not in st.session_state:
    st.session_state.cost_safe_mode = False
if "cost_safe_cap" not in st.session_state:
    st.session_state.cost_safe_cap = 500
if "health_key_status" not in st.session_state:
    st.session_state.health_key_status = "Not tested"
if "health_quota_status" not in st.session_state:
    st.session_state.health_quota_status = "Unknown"
if "health_model_status" not in st.session_state:
    st.session_state.health_model_status = "Unknown"
if "health_last_message" not in st.session_state:
    st.session_state.health_last_message = ""
if "ui_theme_mode" not in st.session_state:
    st.session_state.ui_theme_mode = "pro"
if "settings_loaded" not in st.session_state:
    loaded = load_local_settings()
    st.session_state.api_provider = loaded.get("api_provider", st.session_state.api_provider)
    st.session_state.api_model_choice = loaded.get("api_model_choice", st.session_state.api_model_choice)
    st.session_state.api_model_custom = loaded.get("api_model_custom", st.session_state.api_model_custom)
    st.session_state.cost_safe_mode = loaded.get("cost_safe_mode", st.session_state.cost_safe_mode)
    st.session_state.cost_safe_cap = int(loaded.get("cost_safe_cap", st.session_state.cost_safe_cap))
    st.session_state.ui_theme_mode = loaded.get("ui_theme_mode", st.session_state.ui_theme_mode)
    st.session_state.suno_base_url = loaded.get("suno_base_url", "https://api.sunoapi.org")
    st.session_state.suno_create_path = loaded.get("suno_create_path", "/api/v1/generate")
    st.session_state.suno_status_path = loaded.get("suno_status_path", "/api/v1/generate/record-info?taskId={id}")
    st.session_state.suno_model = loaded.get("suno_model", "chirp-v3-5")
    st.session_state.suno_make_instrumental = bool(loaded.get("suno_make_instrumental", False))
    if st.session_state.api_provider == "Gemini" and st.session_state.api_model_choice not in ["gemini-2.5-flash", "gemini-2.5-pro-latest"]:
        st.session_state.api_model_choice = "gemini-2.5-flash"
    st.session_state.settings_loaded = True
if "lyrics_history" not in st.session_state:
    st.session_state.lyrics_history = []
if "compare_a" not in st.session_state:
    st.session_state.compare_a = ""
if "compare_b" not in st.session_state:
    st.session_state.compare_b = ""
if "compare_c" not in st.session_state:
    st.session_state.compare_c = ""
if "compare_enabled" not in st.session_state:
    st.session_state.compare_enabled = False
if "compare_judge_report" not in st.session_state:
    st.session_state.compare_judge_report = ""
if "compare_judge_winner" not in st.session_state:
    st.session_state.compare_judge_winner = ""
if "compare_judge_profile" not in st.session_state:
    st.session_state.compare_judge_profile = "Strict Commercial"
if "rewrite_output" not in st.session_state:
    st.session_state.rewrite_output = ""
if "section_rewrite_output" not in st.session_state:
    st.session_state.section_rewrite_output = ""
if "suno_generation_id" not in st.session_state:
    st.session_state.suno_generation_id = ""
if "suno_status" not in st.session_state:
    st.session_state.suno_status = ""
if "suno_audio_url" not in st.session_state:
    st.session_state.suno_audio_url = ""
if "suno_last_raw" not in st.session_state:
    st.session_state.suno_last_raw = ""
if "suno_auto_polling" not in st.session_state:
    st.session_state.suno_auto_polling = False
if "suno_poll_interval_sec" not in st.session_state:
    st.session_state.suno_poll_interval_sec = 5
if "hook_last_report" not in st.session_state:
    st.session_state.hook_last_report = {}
if "rhyme_last_report" not in st.session_state:
    st.session_state.rhyme_last_report = {}
if "quick_action" not in st.session_state:
    st.session_state.quick_action = ""
if "active_menu" not in st.session_state:
    st.session_state.active_menu = "compose"

try:
    nav_param = st.query_params.get("nav", "")
except Exception:
    nav_param = ""

if isinstance(nav_param, list):
    nav_param = nav_param[0] if nav_param else ""

if str(nav_param).strip().lower() == "home":
    st.session_state.active_menu = "compose"
    st.session_state.quick_action = ""
    try:
        st.query_params.clear()
    except Exception:
        pass

ensure_daily_usage_state()

if st.session_state.get("ui_theme_mode", "pro") == "classic":
    st.markdown(
        """
<style>
.stApp {
    background: #0b1119 !important;
}
.hero,
.pro-card,
.quickbar,
.block-container {
    backdrop-filter: none !important;
    -webkit-backdrop-filter: none !important;
}
.logo-7code,
.logo-7code::before,
.stButton > button {
    animation: none !important;
}
</style>
""",
        unsafe_allow_html=True,
    )

st.markdown(
    """
<div class="hero">
    <div class="hero-brand">
        <a class="logo-link" href="?nav=home" target="_self" rel="noopener noreferrer"><div class="logo-7code">7 CODE</div></a>
        <div class="hero-title">SEVENZ Studio Pro</div>
    </div>
    <div class="hero-sub">Compose smarter, judge faster, produce cleaner tracks in one professional workflow.</div>
</div>
""",
    unsafe_allow_html=True,
)

menu_col1, menu_col2, menu_col3, menu_col4 = st.columns(4)
with menu_col1:
    if st.button(
        "Compose",
        use_container_width=True,
        key="hero_menu_compose",
        type="primary" if st.session_state.active_menu == "compose" else "secondary",
    ):
        st.session_state.active_menu = "compose"
        st.session_state.quick_action = "generate"
        st.toast("Compose mode ready")
with menu_col2:
    if st.button(
        "Analyze",
        use_container_width=True,
        key="hero_menu_analyze",
        type="primary" if st.session_state.active_menu == "analyze" else "secondary",
    ):
        st.session_state.active_menu = "analyze"
        st.session_state.quick_action = "judge"
        st.toast("Analyze tools ready")
with menu_col3:
    if st.button(
        "Generate",
        use_container_width=True,
        key="hero_menu_generate",
        type="primary" if st.session_state.active_menu == "generate" else "secondary",
    ):
        st.session_state.active_menu = "generate"
        st.session_state.quick_action = "generate"
        st.toast("Generation flow selected")
with menu_col4:
    if st.button(
        "Export",
        use_container_width=True,
        key="hero_menu_export",
        type="primary" if st.session_state.active_menu == "export" else "secondary",
    ):
        st.session_state.active_menu = "export"
        st.session_state.quick_action = "export"
        st.toast("Jumping to export flow")

st.markdown("<div class='quickbar'>", unsafe_allow_html=True)
qa1, qa2, qa3, qa4 = st.columns(4)
with qa1:
    if st.button("🚀 Generate", use_container_width=True, key="quick_generate"):
        st.session_state.active_menu = "generate"
        st.session_state.quick_action = "generate"
with qa2:
    if st.button("📊 Judge", use_container_width=True, key="quick_judge"):
        st.session_state.active_menu = "analyze"
        st.session_state.quick_action = "judge"
with qa3:
    if st.button("💾 Save", use_container_width=True, key="quick_save"):
        st.session_state.active_menu = "compose"
        st.session_state.quick_action = "save"
with qa4:
    if st.button("⬇️ Export TXT", use_container_width=True, key="quick_export"):
        st.session_state.active_menu = "export"
        st.session_state.quick_action = "export"
st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div class='pro-card'><div class='card-kicker'>System Status</div>", unsafe_allow_html=True)
s1, s2, s3, s4, s5 = st.columns(5)
s1.metric("Provider", st.session_state.get("api_provider", "-"))
s2.metric("Model", st.session_state.get("api_model_choice", "-"))
s3.metric("Suno Host", (st.session_state.get("suno_base_url", "-") or "-").replace("https://", "")[:28])
s4.metric("Theme", st.session_state.get("ui_theme_mode", "pro").capitalize())
s5.metric("Version", os.getenv("APP_VERSION", "local-dev"))

if st.session_state.get("health_last_message"):
    st.caption(f"Last error/status: {st.session_state.health_last_message}")

tool_a, tool_b, tool_c = st.columns([1, 1, 1.5])
with tool_a:
    if st.button("🎨 Reset UI Theme", use_container_width=True, key="reset_ui_theme_btn"):
        st.session_state.ui_theme_mode = "classic"
        st.session_state.active_menu = "compose"
        st.session_state.quick_action = ""
        save_local_settings(
            {
                "api_provider": st.session_state.get("api_provider", "Gemini"),
                "api_model_choice": st.session_state.get("api_model_choice", "gemini-2.5-flash"),
                "api_model_custom": st.session_state.get("api_model_custom", ""),
                "cost_safe_mode": bool(st.session_state.get("cost_safe_mode", False)),
                "cost_safe_cap": int(st.session_state.get("cost_safe_cap", 500)),
                "ui_theme_mode": st.session_state.get("ui_theme_mode", "classic"),
                "suno_base_url": st.session_state.get("suno_base_url", "https://api.sunoapi.org"),
                "suno_create_path": st.session_state.get("suno_create_path", "/api/v1/generate"),
                "suno_status_path": st.session_state.get("suno_status_path", "/api/v1/generate/record-info?taskId={id}"),
                "suno_model": st.session_state.get("suno_model", "chirp-v3-5"),
                "suno_make_instrumental": bool(st.session_state.get("suno_make_instrumental", False)),
            }
        )
        st.toast("UI theme reset to classic", icon="✅")
        st.rerun()

settings_snapshot = {
    "api_provider": st.session_state.get("api_provider", "Gemini"),
    "api_model_choice": st.session_state.get("api_model_choice", "gemini-2.5-flash"),
    "api_model_custom": st.session_state.get("api_model_custom", ""),
    "cost_safe_mode": bool(st.session_state.get("cost_safe_mode", False)),
    "cost_safe_cap": int(st.session_state.get("cost_safe_cap", 500)),
    "ui_theme_mode": st.session_state.get("ui_theme_mode", "pro"),
    "suno_base_url": st.session_state.get("suno_base_url", "https://api.sunoapi.org"),
    "suno_create_path": st.session_state.get("suno_create_path", "/api/v1/generate"),
    "suno_status_path": st.session_state.get("suno_status_path", "/api/v1/generate/record-info?taskId={id}"),
    "suno_model": st.session_state.get("suno_model", "chirp-v3-5"),
    "suno_make_instrumental": bool(st.session_state.get("suno_make_instrumental", False)),
}

with tool_b:
    st.download_button(
        "⬇️ Backup Settings",
        data=json.dumps(settings_snapshot, ensure_ascii=False, indent=2),
        file_name="sevenz_settings_backup.json",
        mime="application/json",
        use_container_width=True,
        key="backup_settings_btn",
    )

with tool_c:
    import_file = st.file_uploader("Import settings (.json)", type=["json"], key="settings_import_file")
    if st.button("📥 Apply Imported Settings", use_container_width=True, key="apply_import_settings_btn"):
        if not import_file:
            st.warning("Please choose a JSON settings file first")
        else:
            try:
                imported = json.loads(import_file.getvalue().decode("utf-8"))
                if not isinstance(imported, dict):
                    raise ValueError("Settings file must be a JSON object")

                allowed_keys = {
                    "api_provider",
                    "api_model_choice",
                    "api_model_custom",
                    "cost_safe_mode",
                    "cost_safe_cap",
                    "ui_theme_mode",
                    "suno_base_url",
                    "suno_create_path",
                    "suno_status_path",
                    "suno_model",
                    "suno_make_instrumental",
                }
                imported_clean = {k: v for k, v in imported.items() if k in allowed_keys}
                if not imported_clean:
                    st.warning("No valid settings keys found in file")
                else:
                    merged = load_local_settings()
                    merged.update(imported_clean)
                    save_local_settings(merged)
                    for k, v in imported_clean.items():
                        st.session_state[k] = v
                    st.toast("Settings imported successfully", icon="✅")
                    st.rerun()
            except Exception as exc:
                st.error(f"Failed to import settings: {str(exc)[:140]}")

st.markdown("</div>", unsafe_allow_html=True)

mode = st.radio("⚙️ របៀបប្រើប្រាស់:", ["🆓 ឥតគិតថ្លៃ (Prompt)", "🚀 ស្វ័យប្រវត្តិ (API)"], horizontal=True)

provider = "OpenAI"
model_name = "gpt-4.1-mini"
api_key = ""

if "ស្វ័យប្រវត្តិ" in mode:
    p_col, m_col, t_col = st.columns([1.1, 1.2, 1])

    with p_col:
        provider = st.selectbox("🔌 AI Provider:", PROVIDERS, key="api_provider")

    with m_col:
        model_choices = MODEL_OPTIONS[provider] + ["Custom"]

        if st.session_state.api_model_choice not in model_choices:
            st.session_state.api_model_choice = MODEL_OPTIONS[provider][0]

        model_choice = st.selectbox("🤖 Model:", model_choices, key="api_model_choice")

        if model_choice == "Custom":
            model_name = st.text_input("Custom Model Name", value=st.session_state.api_model_custom, key="api_model_custom")
        else:
            model_name = model_choice

    with t_col:
        key_label = "🔑 Gemini API Key:" if provider == "Gemini" else "🔑 OpenAI API Key:"
        api_key = st.text_input(key_label, type="password")

        resolved_api_key, key_source = resolve_api_key(provider, api_key)
        if resolved_api_key and not api_key.strip():
            source_label = {
                "secrets": "Streamlit secrets",
                "env": "environment variable",
                "dotenv": ".env file",
            }.get(key_source, key_source)
            st.caption(f"✅ Auto-loaded key from {source_label}")
        elif not resolved_api_key:
            st.caption("Tip: ដាក់ key នៅ .env ឬ environment variable ដើម្បីប្រើស្វ័យប្រវត្តិ")

        st.caption("Note: App មិនអាចបង្កើត API key ដោយស្វ័យប្រវត្តិបានទេ, ត្រូវយក key free-tier ពី provider របស់អ្នក។")
        api_key = resolved_api_key

        if st.button("🔍 Test API Key", use_container_width=True):
            if not api_key.strip():
                st.session_state.api_test_message = "status-warn|សូមបញ្ចូល API Key មុនសិន"
            else:
                with st.spinner("Testing API connection..."):
                    ok, result = test_api_key(api_key, provider, model_name)
                if ok and is_api_test_success_response(result):
                    st.session_state.api_test_message = "status-good|API Key ត្រឹមត្រូវ និងដំណើរការ"
                    update_health_from_result(True, result, provider, model_name)
                elif ok:
                    preview = (result or "")[:80].replace("\n", " ").strip()
                    st.session_state.api_test_message = (
                        f"status-warn|Connection បាន ប៉ុន្តែ response ខុស format (ទទួលបាន: {preview})"
                    )
                    update_health_from_result(True, result, provider, model_name)
                elif is_quota_error(result):
                    st.session_state.api_test_message = "status-good|API Key ត្រឹមត្រូវ ប៉ុន្តែ quota អស់ (429). សូមប្ដូរ model ឬគម្រោង billing។"
                    update_health_from_result(False, result, provider, model_name)
                else:
                    st.session_state.api_test_message = f"status-warn|{result}"
                    update_health_from_result(False, result, provider, model_name)

    if st.session_state.api_test_message:
        cls, msg = st.session_state.api_test_message.split("|", 1)
        st.markdown(f"<div class='{cls}'>{msg}</div>", unsafe_allow_html=True)

    tool_col1, tool_col2, tool_col3, tool_col4 = st.columns([1, 1, 1, 1])
    with tool_col1:
        st.session_state.cost_safe_mode = st.checkbox(
            "🛡️ Cost-safe mode",
            value=st.session_state.cost_safe_mode,
            help="Limit max tokens to reduce cost/quota usage.",
        )
    with tool_col2:
        st.session_state.cost_safe_cap = st.number_input(
            "Max Tokens Cap",
            min_value=100,
            max_value=2000,
            value=int(st.session_state.cost_safe_cap),
            step=50,
            disabled=not st.session_state.cost_safe_mode,
        )
    with tool_col3:
        if st.button("💾 Save API Settings", use_container_width=True):
            saved = save_local_settings(
                {
                    "api_provider": provider,
                    "api_model_choice": st.session_state.api_model_choice,
                    "api_model_custom": st.session_state.api_model_custom,
                    "cost_safe_mode": st.session_state.cost_safe_mode,
                    "cost_safe_cap": int(st.session_state.cost_safe_cap),
                    "ui_theme_mode": st.session_state.get("ui_theme_mode", "pro"),
                    "suno_base_url": st.session_state.get("suno_base_url", "https://api.sunoapi.org"),
                    "suno_create_path": st.session_state.get("suno_create_path", "/api/v1/generate"),
                    "suno_status_path": st.session_state.get("suno_status_path", "/api/v1/generate/record-info?taskId={id}"),
                    "suno_model": st.session_state.get("suno_model", "chirp-v3-5"),
                    "suno_make_instrumental": bool(st.session_state.get("suno_make_instrumental", False)),
                }
            )
            st.toast("បានរក្សាទុក settings រួចរាល់" if saved else "រក្សាទុក settings មិនបាន", icon="✅" if saved else "⚠️")

    with tool_col4:
        if st.button("🔄 Reload .env", use_container_width=True):
            st.session_state.dotenv_cache = parse_dotenv_file(ENV_PATH)
            if st.session_state.dotenv_cache:
                st.toast("បាន Reload .env រួចរាល់", icon="✅")
            else:
                st.toast("មិនឃើញ key នៅក្នុង .env ឬ file មិនមាន", icon="⚠️")

    with st.expander("🩺 Health Check & Daily Usage", expanded=True):
        h1, h2, h3, h4 = st.columns(4)
        h1.metric("Key Status", st.session_state.health_key_status)
        h2.metric("Quota", st.session_state.health_quota_status)
        h3.metric("Model", st.session_state.health_model_status)
        h4.metric("Today Requests", str(st.session_state.usage_requests))

        u1, u2, u3 = st.columns(3)
        u1.metric("Prompt Tokens", str(st.session_state.usage_prompt_tokens))
        u2.metric("Completion Tokens", str(st.session_state.usage_completion_tokens))
        u3.metric("Total Tokens", str(st.session_state.usage_total_tokens))

        if st.session_state.health_last_message:
            st.caption(f"Last status: {st.session_state.health_last_message}")

    save_local_settings(
        {
            "api_provider": provider,
            "api_model_choice": st.session_state.api_model_choice,
            "api_model_custom": st.session_state.api_model_custom,
            "cost_safe_mode": st.session_state.cost_safe_mode,
            "cost_safe_cap": int(st.session_state.cost_safe_cap),
            "ui_theme_mode": st.session_state.get("ui_theme_mode", "pro"),
            "suno_base_url": st.session_state.get("suno_base_url", "https://api.sunoapi.org"),
            "suno_create_path": st.session_state.get("suno_create_path", "/api/v1/generate"),
            "suno_status_path": st.session_state.get("suno_status_path", "/api/v1/generate/record-info?taskId={id}"),
            "suno_model": st.session_state.get("suno_model", "chirp-v3-5"),
            "suno_make_instrumental": bool(st.session_state.get("suno_make_instrumental", False)),
        }
    )

if HAS_AUDIO:
    st.caption("✅ Audio recorder module ready")
else:
    st.caption("ℹ️ audio_recorder_streamlit not installed (optional)")


tab1, tab2, tab3, tab4 = st.tabs(
    [
        "1) ✍️ Compose",
        "2) 🧠 Improve",
        "3) 🎚️ Produce",
        "4) 🚀 Visuals + Suno",
    ]
)

with tab1:
    col1, col2, col3 = st.columns(3)
    with col1:
        sel_g = st.selectbox("🎸 Genre:", GENRES)
        final_g = st.text_input("បញ្ចូល Genre:", "") if sel_g == "✍️ វាយបញ្ចូលខ្លួនឯង" else sel_g
        title = st.text_input("🏷️ ចំណងជើង:", "")
    with col2:
        sel_m = st.selectbox("🎭 Mood:", MOODS)
        final_m = st.text_input("បញ្ចូល Mood:", "") if sel_m == "✍️ វាយបញ្ចូលខ្លួនឯង" else sel_m
        bpm = st.selectbox("⏱️ BPM:", ["មិនកំណត់", "70", "100", "120", "150"])
    with col3:
        sel_i = st.selectbox("🎻 Instruments:", INSTS)
        final_i = st.text_input("បញ្ចូល Insts:", "") if sel_i == "✍️ វាយបញ្ចូលខ្លួនឯង" else sel_i
        lang = st.selectbox("🌐 ភាសា:", ["ខ្មែរ", "English", "ចិន (Chinese)"], key="song_lang")

    st.session_state.selected_lang = lang

    desc = st.text_area("📜 សាច់រឿងសង្ខេប:")
    col_a, col_c = st.columns(2)
    struct_name = col_a.selectbox("📊 រចនាបថ:", list(STRUCTURES.keys()))
    preset_choice = col_c.selectbox("🎙️ Presets:", list(STYLES.keys()))
    style_inf = st.text_input(
        "🎙️ Style Influence:",
        value=STYLES[preset_choice] if preset_choice != "ជ្រើសរើស..." else "",
    )

    style_prompt = build_style_summary(final_g, final_m, final_i, style_inf, bpm)
    song_prompt = build_song_prompt(
        lang,
        title,
        desc,
        final_g,
        final_m,
        final_i,
        STRUCTURES[struct_name],
        style_inf,
        bpm,
    )

    st.session_state.compare_enabled = st.checkbox(
        "🆚 Compare Mode (A/B/C)",
        value=st.session_state.compare_enabled,
        help="Generate three versions and select the one you prefer.",
    )

    if st.session_state.compare_enabled:
        st.selectbox(
            "⚖️ AI Judge Style",
            list(JUDGE_PROFILES.keys()),
            index=list(JUDGE_PROFILES.keys()).index(st.session_state.compare_judge_profile)
            if st.session_state.compare_judge_profile in JUDGE_PROFILES
            else 0,
            key="compare_judge_profile",
        )

    generate_clicked = st.button("🎵 បង្កើតទំនុកច្រៀង", use_container_width=True, key="generate_lyrics") or st.session_state.get("quick_action") == "generate"
    if generate_clicked:
        if st.session_state.get("quick_action") == "generate":
            st.session_state.quick_action = ""
        if "ឥតគិតថ្លៃ" in mode:
            st.subheader("🎼 Style Prompt")
            st.code(style_prompt)
            st.subheader("📝 AI Prompt")
            st.code(song_prompt)
        else:
            if not api_key.strip():
                st.error("សូមបញ្ចូល API Key មុនសិន")
            elif not model_name.strip():
                st.error("សូមបញ្ជាក់ Model មុនសិន")
            else:
                if st.session_state.compare_enabled:
                    with st.spinner("Writing compare versions (A/B/C)..."):
                        ok_a, lyrics_a = call_ai(
                            api_key=api_key,
                            system_prompt="You are an expert songwriter. Version A should be cohesive and emotional.",
                            user_prompt=song_prompt,
                            provider=provider,
                            model=model_name,
                            temperature=0.72,
                            max_tokens=get_effective_max_tokens(1200),
                        )
                        ok_b, lyrics_b = call_ai(
                            api_key=api_key,
                            system_prompt="You are an expert songwriter. Version B should have a stronger hook and bolder phrasing.",
                            user_prompt=song_prompt,
                            provider=provider,
                            model=model_name,
                            temperature=0.92,
                            max_tokens=get_effective_max_tokens(1200),
                        )
                        ok_c, lyrics_c = call_ai(
                            api_key=api_key,
                            system_prompt="You are an expert songwriter. Version C should be minimalist, modern, and highly singable with concise lines.",
                            user_prompt=song_prompt,
                            provider=provider,
                            model=model_name,
                            temperature=0.62,
                            max_tokens=get_effective_max_tokens(1200),
                        )

                    if ok_a and ok_b and ok_c:
                        st.session_state.compare_a = lyrics_a
                        st.session_state.compare_b = lyrics_b
                        st.session_state.compare_c = lyrics_c

                        left, mid, right = st.columns(3)
                        with left:
                            st.text_area("Version A", lyrics_a, height=360, key="compare_a_preview")
                            if st.button("✅ Use A", use_container_width=True, key="use_compare_a"):
                                st.session_state.last_lyrics = lyrics_a
                                st.session_state.notepad_content = lyrics_a
                                push_history(lyrics_a, "compare-A", provider, model_name)
                                st.success("បានជ្រើស Version A")
                        with mid:
                            st.text_area("Version B", lyrics_b, height=360, key="compare_b_preview")
                            if st.button("✅ Use B", use_container_width=True, key="use_compare_b"):
                                st.session_state.last_lyrics = lyrics_b
                                st.session_state.notepad_content = lyrics_b
                                push_history(lyrics_b, "compare-B", provider, model_name)
                                st.success("បានជ្រើស Version B")
                        with right:
                            st.text_area("Version C", lyrics_c, height=360, key="compare_c_preview")
                            if st.button("✅ Use C", use_container_width=True, key="use_compare_c"):
                                st.session_state.last_lyrics = lyrics_c
                                st.session_state.notepad_content = lyrics_c
                                push_history(lyrics_c, "compare-C", provider, model_name)
                                st.success("បានជ្រើស Version C")

                        if st.button("⚖️ Judge A/B/C with AI", use_container_width=True, key="judge_compare_abc"):
                            with st.spinner("Judging versions..."):
                                ok_judge, judge_report = judge_lyrics_versions(
                                    api_key,
                                    lyrics_a,
                                    lyrics_b,
                                    lyrics_c,
                                    provider,
                                    model_name,
                                    st.session_state.compare_judge_profile,
                                )

                            if ok_judge:
                                st.session_state.compare_judge_report = judge_report
                                winner_match = re.search(r"WINNER:\s*([ABC])", judge_report, flags=re.IGNORECASE)
                                st.session_state.compare_judge_winner = winner_match.group(1).upper() if winner_match else ""
                            else:
                                st.session_state.compare_judge_report = ""
                                st.session_state.compare_judge_winner = ""
                                update_health_from_result(False, judge_report, provider, model_name)
                                st.error(judge_report)

                        if st.session_state.compare_judge_report:
                            st.subheader("AI Judge Result")
                            st.code(st.session_state.compare_judge_report)

                            winner = st.session_state.compare_judge_winner
                            if winner in ["A", "B", "C"]:
                                st.success(f"AI Winner: Version {winner}")
                                if st.button("🏆 Use AI Winner", use_container_width=True, key="use_ai_winner"):
                                    winner_text = {
                                        "A": lyrics_a,
                                        "B": lyrics_b,
                                        "C": lyrics_c,
                                    }[winner]
                                    st.session_state.last_lyrics = winner_text
                                    st.session_state.notepad_content = winner_text
                                    push_history(winner_text, f"compare-winner-{winner}", provider, model_name)
                                    st.success(f"បានយក Version {winner} តាម AI Judge")
                    else:
                        failed_msg = lyrics_a if not ok_a else lyrics_b if not ok_b else lyrics_c
                        update_health_from_result(False, failed_msg, provider, model_name)
                        st.error(failed_msg)
                else:
                    with st.spinner("Writing lyrics..."):
                        ok, lyrics = call_ai(
                            api_key=api_key,
                            system_prompt="You are an expert songwriter. Write complete, emotional, high-quality lyrics.",
                            user_prompt=song_prompt,
                            provider=provider,
                            model=model_name,
                            temperature=0.8,
                            max_tokens=get_effective_max_tokens(1200),
                        )

                    if ok:
                        generated_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        st.session_state.last_lyrics = lyrics
                        st.session_state.notepad_content = (
                            f"TITLE\n{title or 'Untitled'}\n\n"
                            f"GENERATED AT\n{generated_at}\n\n"
                            f"PROVIDER\n{provider}\n\n"
                            f"MODEL\n{model_name}\n\n"
                            f"STYLE\n{style_prompt}\n\n"
                            f"LYRICS\n\n{lyrics}"
                        )
                        push_history(lyrics, "generate", provider, model_name)
                        st.success("Lyrics generated successfully!")
                        st.text_area("Lyrics", lyrics, height=500)
                        st.caption(f"Word count: {len(lyrics.split())}")
                    else:
                        update_health_from_result(False, lyrics, provider, model_name)
                        st.error(lyrics)

with tab2:
    st.header("👤 កែសម្រួល & វាយតម្លៃទំនុកច្រៀង")
    ai_text = st.text_area("Paste ទំនុកច្រៀង:", value=st.session_state.last_lyrics, height=260)
    if not ai_text.strip():
        st.caption("Start by pasting lyrics. បន្ទាប់មកប្រើ Analysis tools ខាងក្រោម។")

    st.markdown("<div class='pro-card'><div class='card-kicker'>Analysis</div>", unsafe_allow_html=True)
    with st.expander("🎯 Hook Score + Singability Checker", expanded=True):
        if st.button("Check Hook/Singability", use_container_width=True, key="hook_singability_check"):
            report = hook_and_singability_report(ai_text)
            st.session_state.hook_last_report = report

        report = st.session_state.get("hook_last_report", {})
        if report:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Hook Score", str(report.get("hook_score", 0)))
            c2.metric("Singability", str(report.get("singability_score", 0)))
            c3.metric("Avg Words/Line", str(report.get("avg_words", 0)))
            c4.metric("Avg Syllables/Line", str(report.get("avg_syllables", 0)))
            render_score_bar("Hook", report.get("hook_score", 0))
            render_score_bar("Singability", report.get("singability_score", 0))
            if report.get("suggestions"):
                st.markdown("\n".join(f"- {s}" for s in report["suggestions"]))

    with st.expander("🎵 Rhyme Engine (Khmer/English basic)", expanded=False):
        if st.button("Analyze Rhyme", use_container_width=True, key="rhyme_engine_check"):
            rhyme = rhyme_engine_report(ai_text)
            st.session_state.rhyme_last_report = rhyme

        rhyme = st.session_state.get("rhyme_last_report", {})
        if rhyme:
            st.metric("Rhyme Score", str(rhyme.get("rhyme_score", 0)))
            render_score_bar("Rhyme", rhyme.get("rhyme_score", 0))
            st.caption(f"Scheme: {rhyme.get('scheme', 'N/A')}")
            if rhyme.get("pairs"):
                st.markdown("\n".join(rhyme["pairs"]))
            if rhyme.get("suggestions"):
                st.markdown("\n".join(f"- {s}" for s in rhyme["suggestions"]))
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='pro-card'><div class='card-kicker'>Rewrite Lab</div>", unsafe_allow_html=True)
    rw_col1, rw_col2 = st.columns([1.4, 1])
    with rw_col1:
        rewrite_action = st.selectbox("✨ Smart Rewrite Action", list(REWRITE_ACTIONS.keys()))
    with rw_col2:
        if st.button("🔁 Rewrite with AI", use_container_width=True, key="smart_rewrite_btn"):
            if not ai_text.strip():
                st.warning("សូមបញ្ចូលទំនុកច្រៀងជាមុនសិន")
            elif "ឥតគិតថ្លៃ" in mode:
                st.info("Free mode: សូមប្រើ API mode សម្រាប់ Smart Rewrite")
            elif not api_key.strip():
                st.error("Please enter your API Key.")
            else:
                with st.spinner("Rewriting lyrics..."):
                    ok_rw, rw_text = smart_rewrite(api_key, ai_text, REWRITE_ACTIONS[rewrite_action], provider, model_name)
                if ok_rw:
                    st.session_state.rewrite_output = rw_text
                    st.text_area("Rewrite Output", rw_text, height=240, key="rewrite_output_view")
                    if st.button("✅ Use Rewrite", use_container_width=True, key="use_rewrite"):
                        st.session_state.last_lyrics = rw_text
                        st.session_state.notepad_content = rw_text
                        push_history(rw_text, f"rewrite:{rewrite_action}", provider, model_name)
                        st.success("បានយក Rewrite output")
                else:
                    update_health_from_result(False, rw_text, provider, model_name)
                    st.error(rw_text)

    if st.session_state.rewrite_output:
        st.text_area("Last Rewrite", st.session_state.rewrite_output, height=180, key="rewrite_last_view")

    with st.expander("✂️ Section-by-Section Rewrite", expanded=False):
        parsed_sections = parse_lyric_sections(ai_text)
        section_options = ["Full Lyrics", "Auto (First Chorus/Hook)"] + [hdr for hdr, _ in parsed_sections]
        selected_section = st.selectbox("Target Section", section_options, key="target_rewrite_section")
        section_instruction = st.text_input(
            "Rewrite instruction",
            value="Make this section catchier with tighter rhyme and simpler words.",
            key="target_rewrite_instruction",
        )

        if st.button("Rewrite Selected Section", use_container_width=True, key="section_rewrite_btn"):
            if not ai_text.strip():
                st.warning("សូមបញ្ចូលទំនុកច្រៀងជាមុនសិន")
            elif "ឥតគិតថ្លៃ" in mode:
                st.info("Free mode: សូមប្រើ API mode សម្រាប់ Section Rewrite")
            elif not api_key.strip():
                st.error("Please enter your API Key.")
            else:
                with st.spinner("Rewriting selected section..."):
                    ok_sec, sec_text = rewrite_section_with_ai(
                        api_key,
                        ai_text,
                        selected_section,
                        section_instruction,
                        provider,
                        model_name,
                    )
                if ok_sec:
                    st.session_state.section_rewrite_output = sec_text
                    st.text_area("Section Rewrite Output", sec_text, height=280, key="section_rewrite_output_view")
                    if st.button("✅ Use Section Rewrite", use_container_width=True, key="use_section_rewrite"):
                        st.session_state.last_lyrics = sec_text
                        st.session_state.notepad_content = sec_text
                        push_history(sec_text, f"section-rewrite:{selected_section}", provider, model_name)
                        st.success("បានយកលទ្ធផល Section Rewrite")
                else:
                    update_health_from_result(False, sec_text, provider, model_name)
                    st.error(sec_text)

        if st.session_state.section_rewrite_output:
            st.text_area("Last Section Rewrite", st.session_state.section_rewrite_output, height=180, key="section_rewrite_last_view")
    st.markdown("</div>", unsafe_allow_html=True)

    evaluate_clicked = st.button("📊 វាយតម្លៃគុណភាពទំនុកច្រៀង", key="evaluate_lyrics") or st.session_state.get("quick_action") == "judge"
    if evaluate_clicked:
        if st.session_state.get("quick_action") == "judge":
            st.session_state.quick_action = ""
        if not ai_text.strip():
            st.warning("សូមបញ្ចូលទំនុកច្រៀងជាមុនសិន")
        elif "ឥតគិតថ្លៃ" in mode:
            st.info("Free mode: សូមប្រើ API mode ដើម្បីវាយតម្លៃដោយ AI")
        elif not api_key.strip():
            st.error("Please enter your API Key.")
        else:
            with st.spinner("Analyzing lyrics..."):
                ok, result = analyze_lyrics(api_key, ai_text, provider, model_name)
            if ok:
                st.markdown(result)
            else:
                update_health_from_result(False, result, provider, model_name)
                st.error(result)

    rewritten = st.text_area("🛠️ កែសម្រួលដោយខ្លួនឯង:", value=ai_text, height=220)
    save_clicked = st.button("💾 រក្សាទុកទៅ Notepad", key="save_edit") or st.session_state.get("quick_action") == "save"
    if save_clicked:
        if st.session_state.get("quick_action") == "save":
            st.session_state.quick_action = ""
        st.session_state.notepad_content = rewritten
        st.session_state.last_lyrics = rewritten
        push_history(rewritten, "manual-edit", provider, model_name)
        st.success("បានរក្សាទុករួចរាល់")

    with st.expander("🕘 Version History / Undo", expanded=False):
        if not st.session_state.lyrics_history:
            st.caption("មិនទាន់មាន history នៅឡើយ")
        else:
            labels = [
                f"{idx + 1}. {item['time']} | {item['source']} | {item['provider']}:{item['model']}"
                for idx, item in enumerate(st.session_state.lyrics_history)
            ]
            selected_label = st.selectbox("History", labels, key="history_select")
            selected_idx = labels.index(selected_label)
            selected_item = st.session_state.lyrics_history[selected_idx]

            st.text_area("Selected Version", selected_item["content"], height=220, key="history_preview")

            h_col1, h_col2 = st.columns(2)
            with h_col1:
                if st.button("↩️ Restore Selected", use_container_width=True, key="restore_selected"):
                    st.session_state.last_lyrics = selected_item["content"]
                    st.session_state.notepad_content = selected_item["content"]
                    st.success("បាន Restore version ដែលបានជ្រើស")
            with h_col2:
                if st.button("⏪ Undo Last", use_container_width=True, key="undo_last"):
                    if len(st.session_state.lyrics_history) > 1:
                        st.session_state.lyrics_history.pop(0)
                        previous = st.session_state.lyrics_history[0]
                        st.session_state.last_lyrics = previous["content"]
                        st.session_state.notepad_content = previous["content"]
                        st.success("បាន Undo ទៅ version មុន")
                    else:
                        st.info("មិនមាន version មុនសម្រាប់ Undo")

with tab3:
    st.subheader("📓 កន្លែងកត់ត្រា")
    if st.session_state.get("quick_action") == "export":
        st.info("Quick Export: សូមចុចប៊ូតុង TXT/PDF/DOCX ខាងក្រោម")
        st.session_state.quick_action = ""
    st.session_state.notepad_content = st.text_area(
        "",
        value=st.session_state.notepad_content,
        height=320,
        key="main_notepad",
    )

    if any(ord(ch) > 255 for ch in st.session_state.notepad_content):
        st.caption("PDF អាចមិនគាំទ្រអក្សរខ្មែរគ្រប់តួ។ សូមប្រើ DOCX សម្រាប់ Unicode ពេញលេញ។")
    if not HAS_DOCX:
        st.caption("DOCX export មិនទាន់មាន (python-docx not installed in current interpreter)")

    col_tts, col_txt, col_pdf, col_docx = st.columns(4)

    with col_tts:
        if st.button("🔊 អានទំនុកច្រៀង (TTS)", use_container_width=True, key="tts_btn"):
            if not st.session_state.notepad_content.strip():
                st.warning("គ្មានអត្ថបទសម្រាប់អាន")
            else:
                tts = gTTS(
                    text=st.session_state.notepad_content[:700],
                    lang=LANG_TTS.get(st.session_state.selected_lang, "en"),
                )
                audio_buf = BytesIO()
                tts.write_to_fp(audio_buf)
                st.audio(audio_buf.getvalue(), format="audio/mp3")

    with col_txt:
        st.download_button(
            "⬇️ TXT",
            data=st.session_state.notepad_content,
            file_name="sevenz_lyrics.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with col_pdf:
        try:
            pdf_data = make_pdf_bytes("SEVENZ Lyrics Export", st.session_state.notepad_content, provider, model_name)
            st.download_button(
                "⬇️ PDF",
                data=pdf_data,
                file_name="sevenz_lyrics.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as pdf_exc:
            st.button("⬇️ PDF", use_container_width=True, disabled=True)
            st.caption(f"PDF export unavailable right now: {str(pdf_exc)[:120]}")
            if HAS_DOCX:
                fallback_docx = make_docx_bytes(
                    "SEVENZ Lyrics Export",
                    st.session_state.notepad_content,
                    provider,
                    model_name,
                )
                st.download_button(
                    "⬇️ DOCX (Fallback)",
                    data=fallback_docx,
                    file_name="sevenz_lyrics_fallback.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="docx_fallback_btn",
                )

    with col_docx:
        if HAS_DOCX:
            docx_data = make_docx_bytes("SEVENZ Lyrics Export", st.session_state.notepad_content, provider, model_name)
            st.download_button(
                "⬇️ DOCX",
                data=docx_data,
                file_name="sevenz_lyrics.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.button("⬇️ DOCX", disabled=True, use_container_width=True)

with tab4:
    st.markdown("<div class='pro-card'><div class='card-kicker'>Creative Tools</div>", unsafe_allow_html=True)
    st.subheader("🎨 បង្កើត Prompt រូបភាព")
    if st.button("បង្កើត Prompt", key="img_prompt_btn"):
        st.code(
            f"Album cover art for {final_g}, mood {final_m}, based on story: {desc}. "
            "Cinematic lighting, ultra-detailed, 8k."
        )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='pro-card'><div class='card-kicker'>Translator</div>", unsafe_allow_html=True)
    st.subheader("🌐 មុខងារបកប្រែរក្សាពាក្យជួន (Smart Converter)")
    source_text = st.text_area("បញ្ចូលអត្ថបទដើម:", height=220)
    target_lang = st.selectbox("បកប្រែទៅភាសា:", ["ខ្មែរ", "English", "ចិន (Chinese)"], key="target_lang")

    if st.button("បកប្រែដោយរក្សាពាក្យជួន", key="smart_translate"):
        if not source_text.strip():
            st.warning("សូមបញ្ចូលអត្ថបទដើមជាមុនសិន")
        elif "ឥតគិតថ្លៃ" in mode:
            st.info("Free mode: សូមប្រើ API mode សម្រាប់ Smart Translation")
        elif not api_key.strip():
            st.error("Please enter your API Key.")
        else:
            with st.spinner("កំពុងបកប្រែដោយរក្សាពាក្យជួន..."):
                ok, translated = translate_with_rhyme(api_key, source_text, target_lang, provider, model_name)
            if ok:
                st.text_area("លទ្ធផលបកប្រែ", translated, height=280)
                st.session_state.notepad_content = translated
                st.session_state.last_lyrics = translated
                push_history(translated, "translate", provider, model_name)
            else:
                update_health_from_result(False, translated, provider, model_name)
                st.error(translated)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='pro-card'><div class='card-kicker'>Suno Control Center</div>", unsafe_allow_html=True)
    st.subheader("🎧 Direct Music Generation (Suno API)")
    st.caption("Suno endpoints ខុសគ្នាតាម provider/proxy។ អ្នកអាចកំណត់ Base URL និង paths ខាងក្រោម។")

    status_view = (st.session_state.get("suno_status") or "idle").strip().lower()
    st.markdown(f"<span class='pill'>Status: {status_view or 'idle'}</span>", unsafe_allow_html=True)

    if not HAS_REQUESTS:
        st.warning("មិនទាន់មាន requests package។ សូម install មុន: pip install requests")

    suno_key_input = st.text_input("🔑 Suno API Key", type="password", key="suno_api_key_input")
    resolved_suno_key, suno_key_source = resolve_named_api_key(suno_key_input, ENV_KEY_CANDIDATES.get("Suno", []))
    if resolved_suno_key and not suno_key_input.strip():
        source_label = {
            "secrets": "Streamlit secrets",
            "env": "environment variable",
            "dotenv": ".env file",
        }.get(suno_key_source, suno_key_source)
        st.caption(f"✅ Auto-loaded Suno key from {source_label}")

    st.caption("Configure")

    if st.button("Reset to SunoAPI defaults", use_container_width=True, key="suno_reset_defaults_btn"):
        st.session_state.suno_base_url = "https://api.sunoapi.org"
        st.session_state.suno_create_path = "/api/v1/generate"
        st.session_state.suno_status_path = "/api/v1/generate/record-info?taskId={id}"
        st.session_state.suno_model = "V4_5ALL"
        st.session_state.suno_make_instrumental = False
        save_local_settings(
            {
                "api_provider": st.session_state.get("api_provider", "OpenAI"),
                "api_model_choice": st.session_state.get("api_model_choice", "Auto"),
                "api_model_custom": st.session_state.get("api_model_custom", ""),
                "cost_safe_mode": bool(st.session_state.get("cost_safe_mode", False)),
                "cost_safe_cap": int(st.session_state.get("cost_safe_cap", 2000)),
                "ui_theme_mode": st.session_state.get("ui_theme_mode", "pro"),
                "suno_base_url": st.session_state.suno_base_url,
                "suno_create_path": st.session_state.suno_create_path,
                "suno_status_path": st.session_state.suno_status_path,
                "suno_model": st.session_state.suno_model,
                "suno_make_instrumental": st.session_state.suno_make_instrumental,
            }
        )
        st.success("បាន reset ទៅ SunoAPI defaults រួចរាល់")
        st.rerun()

    s_col1, s_col2 = st.columns(2)
    with s_col1:
        st.text_input("Base URL", key="suno_base_url", placeholder="https://api.sunoapi.org")
        st.text_input("Create Path", key="suno_create_path", placeholder="/api/v1/generate")
    with s_col2:
        st.text_input("Status Path Template", key="suno_status_path", placeholder="/api/v1/generate/record-info?taskId={id}")
        st.text_input("Model", key="suno_model", placeholder="chirp-v3-5")

    st.checkbox(
        "Instrumental only",
        value=bool(st.session_state.get("suno_make_instrumental", False)),
        key="suno_make_instrumental",
    )

    st.caption("Generate")
    default_music_prompt = f"{final_g} | {final_m} | {style_inf or 'modern clean production'}"
    suno_prompt = st.text_area("Music Prompt", value=default_music_prompt, height=110, key="suno_prompt")
    suno_title = st.text_input("Track Title", value=title or "SEVENZ Track", key="suno_title")
    suno_lyrics = st.text_area(
        "Lyrics (optional)",
        value=st.session_state.last_lyrics,
        height=180,
        key="suno_lyrics_input",
    )

    b1, b2, b3 = st.columns(3)
    with b1:
        create_clicked = st.button("🚀 Create Suno Generation", use_container_width=True, key="suno_create_btn")
    with b2:
        check_clicked = st.button("🔄 Check Suno Status", use_container_width=True, key="suno_check_btn")
    with b3:
        clear_clicked = st.button("🧹 Clear Suno Result", use_container_width=True, key="suno_clear_btn")

    ap1, ap2 = st.columns([1.2, 1])
    with ap1:
        st.checkbox(
            "Auto polling status",
            value=bool(st.session_state.get("suno_auto_polling", False)),
            key="suno_auto_polling",
        )
    with ap2:
        st.number_input(
            "Poll interval (sec)",
            min_value=2,
            max_value=30,
            value=int(st.session_state.get("suno_poll_interval_sec", 5)),
            step=1,
            key="suno_poll_interval_sec",
        )

    if clear_clicked:
        st.session_state.suno_generation_id = ""
        st.session_state.suno_status = ""
        st.session_state.suno_audio_url = ""
        st.session_state.suno_last_raw = ""
        st.success("បានសម្អាតលទ្ធផល Suno")

    if create_clicked:
        if not HAS_REQUESTS:
            st.error("requests package មិនទាន់មាន")
        elif not resolved_suno_key.strip():
            st.error("សូមបញ្ចូល Suno API key")
        elif not st.session_state.suno_base_url.strip():
            st.error("សូមបញ្ចូល Base URL")
        elif not suno_prompt.strip() and not suno_lyrics.strip():
            st.error("សូមបញ្ចូល Music Prompt ឬ Lyrics")
        else:
            payload = {
                "prompt": suno_prompt.strip(),
                "title": (suno_title or "SEVENZ Track").strip(),
                "lyrics": suno_lyrics.strip(),
                "tags": f"{final_g}, {final_m}",
                "model": (st.session_state.suno_model or "chirp-v3-5").strip(),
                "instrumental": bool(st.session_state.suno_make_instrumental),
            }
            with st.spinner("Creating Suno generation..."):
                ok_create, create_msg, generation_id, raw = suno_create_generation(
                    resolved_suno_key,
                    st.session_state.suno_base_url,
                    st.session_state.suno_create_path,
                    payload,
                )
            st.session_state.suno_last_raw = json.dumps(raw, ensure_ascii=False, indent=2)
            if ok_create:
                st.session_state.suno_generation_id = generation_id or ""
                if generation_id:
                    st.success(f"បានបង្កើត job រួចរាល់: {generation_id}")
                else:
                    st.warning("បានបង្កើត request តែរកមិនឃើញ generation id ពី response")
            else:
                st.error(create_msg)

    st.text_input("Generation ID", key="suno_generation_id", placeholder="paste generation id here")

    if check_clicked:
        if not HAS_REQUESTS:
            st.error("requests package មិនទាន់មាន")
        elif not resolved_suno_key.strip():
            st.error("សូមបញ្ចូល Suno API key")
        elif not st.session_state.suno_base_url.strip():
            st.error("សូមបញ្ចូល Base URL")
        elif not st.session_state.suno_generation_id.strip():
            st.error("សូមបញ្ចូល Generation ID")
        else:
            with st.spinner("Checking Suno status..."):
                ok_status, status_msg, status_text, audio_url, raw = suno_check_generation(
                    resolved_suno_key,
                    st.session_state.suno_base_url,
                    st.session_state.suno_status_path,
                    st.session_state.suno_generation_id.strip(),
                )
            st.session_state.suno_last_raw = json.dumps(raw, ensure_ascii=False, indent=2)
            if ok_status:
                st.session_state.suno_status = status_text or "unknown"
                st.session_state.suno_audio_url = audio_url or ""
                st.success(f"Status: {st.session_state.suno_status}")
            else:
                st.error(status_msg)

    status_lower = (st.session_state.suno_status or "").strip().lower()
    pending_states = {"", "queued", "pending", "processing", "running", "submitted", "in_progress"}
    terminal_states = {"completed", "done", "failed", "error", "canceled", "cancelled"}
    can_auto_poll = (
        bool(st.session_state.get("suno_auto_polling", False))
        and not check_clicked
        and bool(st.session_state.suno_generation_id.strip())
        and not st.session_state.suno_audio_url
        and status_lower in pending_states
    )

    if can_auto_poll:
        if not HAS_REQUESTS:
            st.error("Auto polling ត្រូវការ requests package")
        elif not resolved_suno_key.strip():
            st.warning("Auto polling បានបើក ប៉ុន្តែមិនមាន Suno API key")
        elif not st.session_state.suno_base_url.strip():
            st.warning("Auto polling បានបើក ប៉ុន្តែមិនមាន Base URL")
        else:
            interval_sec = int(st.session_state.get("suno_poll_interval_sec", 5))
            with st.spinner(f"Auto polling... refresh every {interval_sec}s"):
                ok_status, status_msg, status_text, audio_url, raw = suno_check_generation(
                    resolved_suno_key,
                    st.session_state.suno_base_url,
                    st.session_state.suno_status_path,
                    st.session_state.suno_generation_id.strip(),
                )
            st.session_state.suno_last_raw = json.dumps(raw, ensure_ascii=False, indent=2)
            if ok_status:
                st.session_state.suno_status = status_text or "unknown"
                st.session_state.suno_audio_url = audio_url or ""
                next_status = (st.session_state.suno_status or "").strip().lower()
                if st.session_state.suno_audio_url or next_status in terminal_states:
                    st.success(f"Auto polling stopped at status: {st.session_state.suno_status}")
                else:
                    st.info(f"Auto polling status: {st.session_state.suno_status} (next in {interval_sec}s)")
                    time.sleep(interval_sec)
                    st.rerun()
            else:
                st.error(status_msg)

    st.caption("Monitor")
    if st.session_state.suno_status:
        st.caption(f"Current status: {st.session_state.suno_status}")

    if st.session_state.suno_audio_url:
        st.audio(st.session_state.suno_audio_url)
        st.markdown(f"[Open audio link]({st.session_state.suno_audio_url})")

    if st.session_state.suno_last_raw:
        with st.expander("Suno Raw Response", expanded=False):
            st.code(st.session_state.suno_last_raw)
    st.markdown("</div>", unsafe_allow_html=True)

app_version = os.getenv("APP_VERSION", "local-dev")
st.caption(f"SEVENZ Studio | version: {app_version} | updated: {datetime.datetime.now().strftime('%Y-%m-%d')}")
